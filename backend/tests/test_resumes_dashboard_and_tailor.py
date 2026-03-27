"""Integration-style tests for resume ingestion, dashboard aggregation, job match, and tailoring endpoints."""

from io import BytesIO
from bson import ObjectId
from docx import Document


def _create_confirmation_and_evidence(client, headers, user_id, skill_python, skill_ml):
    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [
                {"skill_id": skill_python, "proficiency": 4},
                {"skill_id": skill_ml, "proficiency": 3},
            ],
            "rejected": [],
            "edited": [],
        },
    )
    client.post(
        "/evidence/",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "project",
            "title": "Analytics Project",
            "source": "manual-entry",
            "text_excerpt": "Built Python ML dashboards and APIs.",
            "skill_ids": [skill_python, skill_ml],
            "origin": "user",
        },
    )
    client.post(
        "/evidence/",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "project",
            "title": "Reporting API",
            "source": "manual-entry",
            "text_excerpt": "Implemented reporting APIs, analytics data flows, and stakeholder-ready dashboards.",
            "skill_ids": [skill_python],
            "origin": "user",
        },
    )
    client.post(
        "/portfolio/items",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "project",
            "title": "Analytics Platform",
            "summary": "Built internal analytics systems.",
            "bullets": ["Delivered Python and ML workflows."],
            "skill_ids": [skill_python, skill_ml],
            "visibility": "private",
            "priority": 2,
        },
    )


def test_resume_ingest_promote_dashboard_and_tailor_endpoints(test_context, monkeypatch):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_python = test_context["skill_python"]
    skill_ml = test_context["skill_ml"]

    monkeypatch.setattr("app.routers.resumes.extract_pdf_text", lambda _bytes: "Python ML FastAPI " * 10)

    text_ingest = client.post(
        "/ingest/resume/text",
        headers=headers,
        json={
            "user_id": user_id,
            "text": "\n".join(
                [
                    "Tester Example",
                    "SUMMARY",
                    "Python and ML developer.",
                    "EXPERIENCE",
                    "Capstone Project",
                    "- Built Python ML dashboards.",
                    "- Delivered FastAPI APIs.",
                    "EDUCATION",
                    "BS Computer Science",
                ]
            ),
        },
    )
    assert text_ingest.status_code == 200
    snapshot_id = text_ingest.json()["snapshot_id"]
    assert any(
        str(doc.get("source_id")) == snapshot_id and doc.get("source_type") == "resume_snapshot"
        for doc in test_context["db"]["rag_chunks"].docs
    )

    pdf_ingest = client.post(
        "/ingest/resume/pdf",
        headers=headers,
        files={"file": ("resume.pdf", b"%PDF-1.4 fake", "application/pdf")},
        data={"user_id": user_id},
    )
    assert pdf_ingest.status_code == 200

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={"resume_snapshot_id": snapshot_id, "confirmed": [{"skill_id": skill_python, "proficiency": 4}], "rejected": [], "edited": []},
    )

    promoted = client.post(f"/ingest/resume/{snapshot_id}/promote", headers=headers, data={"user_id": user_id})
    assert promoted.status_code == 200

    _create_confirmation_and_evidence(client, headers, user_id, skill_python, skill_ml)

    summary = client.get("/dashboard/summary", headers=headers)
    assert summary.status_code == 200
    assert summary.json()["totals"]["confirmed_skills"] >= 1
    assert "portfolio_to_job_analytics" in summary.json()
    assert "portfolio_type_distribution" in summary.json()
    assert "recent_match_trend" in summary.json()

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "We need Python, ML, FastAPI, APIs, and analytics experience. Responsibilities include dashboards and model delivery." * 2,
        },
    )
    assert ingest.status_code == 200
    job_id = ingest.json()["id"]

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": job_id})
    assert match.status_code == 200
    history_id = match.json()["history_id"]
    assert "gap_reasoning_summary" in match.json()
    assert "gap_insights" in match.json()
    assert "personal_skill_vector_score" in match.json()
    assert any(item["label"] == "Personal skill vector" for item in match.json()["score_breakdown"])

    user_vector = client.get("/tailor/user-vector", headers=headers)
    assert user_vector.status_code == 200
    assert user_vector.json()["embedding_dimensions"] > 0

    vector_history = client.get("/tailor/user-vector/history", headers=headers)
    assert vector_history.status_code == 200
    assert vector_history.json()

    refreshed_summary = client.get("/dashboard/summary", headers=headers)
    assert refreshed_summary.status_code == 200
    dashboard_payload = refreshed_summary.json()
    assert isinstance(dashboard_payload["portfolio_to_job_analytics"]["matched_skill_rate_pct"], float)
    assert dashboard_payload["portfolio_to_job_analytics"]["matched_skill_rate_pct"] > 0
    assert dashboard_payload["recent_match_trend"]

    # Older saved runs may have counts persisted without the full matched_skill_ids array.
    # The dashboard metric should still be driven by the stored counts in that case.
    job_match_docs = test_context["db"]["job_match_runs"].docs
    assert job_match_docs
    job_match_docs[0]["analysis"]["matched_skill_ids"] = []

    fallback_summary = client.get("/dashboard/summary", headers=headers)
    assert fallback_summary.status_code == 200
    assert fallback_summary.json()["portfolio_to_job_analytics"]["matched_skill_rate_pct"] > 0

    preview = client.post("/tailor/preview", headers=headers, json={"user_id": user_id, "job_id": job_id, "resume_snapshot_id": snapshot_id})
    assert preview.status_code == 200
    tailored_id = preview.json()["id"]
    assert preview.json()["retrieved_context"]

    rag_search = client.get("/tailor/rag/search", headers=headers, params={"q": "Python ML dashboards", "limit": 3})
    assert rag_search.status_code == 200
    assert rag_search.json()

    resumes = client.get("/tailor/resumes", headers=headers, params={"user_id": user_id})
    assert resumes.status_code == 200
    assert any(item["id"] == tailored_id for item in resumes.json())

    resume_detail = client.get(f"/tailor/resumes/{tailored_id}", headers=headers, params={"user_id": user_id})
    assert resume_detail.status_code == 200

    history = client.get("/tailor/history", headers=headers, params={"user_id": user_id})
    assert history.status_code == 200
    assert any(item["id"] == history_id for item in history.json())

    history_detail = client.get(f"/tailor/history/{history_id}", headers=headers, params={"user_id": user_id})
    assert history_detail.status_code == 200

    compare = client.get("/tailor/history/compare", headers=headers, params={"user_id": user_id, "left_id": history_id, "right_id": history_id})
    assert compare.status_code == 200

    settings = client.get("/tailor/settings/status", headers=headers)
    assert settings.status_code == 200
    assert settings.json()["provider_mode"] == "test"

    rewrite = client.post(f"/tailor/{tailored_id}/rewrite", headers=headers, json={"focus": "ats"})
    assert rewrite.status_code == 200

    docx = client.get(f"/tailor/{tailored_id}/export/docx", headers=headers)
    assert docx.status_code == 200

    pdf = client.get(f"/tailor/{tailored_id}/export/pdf", headers=headers)
    assert pdf.status_code == 200

    delete_history = client.delete(f"/tailor/history/{history_id}", headers=headers, params={"user_id": user_id})
    assert delete_history.status_code == 200

    delete_resume = client.delete(f"/tailor/resumes/{tailored_id}", headers=headers, params={"user_id": user_id})
    assert delete_resume.status_code == 200


def test_tailored_resume_uses_user_resume_template_and_ai_preferences(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    db = test_context["db"]

    resume_text = "\n".join(
        [
            "Jane Student",
            "SUMMARY",
            "Builder focused on analytics systems and ML delivery.",
            "SKILLS",
            "Python, SQL, Communication",
            "EXPERIENCE",
            "Capstone Research Assistant",
            "- Built internal dashboards for faculty.",
            "EDUCATION",
            "BS Computer Science",
        ]
    )
    ingest = client.post("/ingest/resume/text", headers=headers, json={"user_id": user_id, "text": resume_text})
    assert ingest.status_code == 200
    snapshot_id = ingest.json()["snapshot_id"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [
                {"skill_id": test_context["skill_python"], "proficiency": 4},
                {"skill_id": test_context["skill_ml"], "proficiency": 3},
            ],
            "rejected": [],
            "edited": [],
        },
    )
    client.post(
        "/evidence/",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "project",
            "title": "ML Dashboard",
            "source": "manual-entry",
            "text_excerpt": "Built Python ML dashboard APIs and shipped analytics views.",
            "skill_ids": [test_context["skill_python"], test_context["skill_ml"]],
            "origin": "user",
        },
    )
    client.post(
        "/evidence/",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "project",
            "title": "Reporting API",
            "source": "manual-entry",
            "text_excerpt": "Implemented reporting APIs, analytics data flows, and stakeholder-ready dashboards.",
            "skill_ids": [test_context["skill_python"]],
            "origin": "user",
        },
    )

    ai_settings = client.get("/tailor/settings/preferences", headers=headers)
    assert ai_settings.status_code == 200
    assert ai_settings.json()["preferences"]["inference_mode"]

    patched = client.patch(
        "/tailor/settings/preferences",
        headers=headers,
        json={
          "inference_mode": "local-fallback",
          "embedding_model": "sentence-transformers/all-MiniLM-L6-v2",
          "zero_shot_model": "MoritzLaurer/deberta-v3-base-zeroshot-v1.1-all-33",
        },
    )
    assert patched.status_code == 200
    assert patched.json()["preferences"]["inference_mode"] == "local-fallback"

    job = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "Analytics Engineer",
            "company": "SkillBridge",
            "location": "Remote",
            "text": "Looking for Python, ML, dashboard, analytics, and API delivery experience for a product analytics role." * 2,
        },
    )
    assert job.status_code == 200
    job_id = job.json()["id"]

    preview = client.post(
        "/tailor/preview",
        headers=headers,
        json={"user_id": user_id, "job_id": job_id, "resume_snapshot_id": snapshot_id, "template": "modern_v1"},
    )
    assert preview.status_code == 200
    payload = preview.json()
    assert payload["resume_snapshot_id"] == snapshot_id
    assert payload["template_source"] == "user_resume"
    assert payload["template"] == "modern_v1"
    assert payload["retrieved_context"]
    assert len(payload["selected_item_ids"]) >= 2
    section_titles = [section["title"] for section in payload["sections"]]
    assert "Profile" in section_titles
    assert "Education" in section_titles
    assert any(section["title"] == "Selected Impact" for section in payload["sections"])
    summary_section = next(section for section in payload["sections"] if section["title"] == "Profile")
    assert any("Analytics Engineer at SkillBridge" in line for line in summary_section["lines"])
    assert any(
        line.startswith("- Rewritten (impact):")
        for section in payload["sections"]
        for line in section["lines"]
    )

    ats_preview = client.post(
        "/tailor/preview",
        headers=headers,
        json={"user_id": user_id, "job_id": job_id, "resume_snapshot_id": snapshot_id, "template": "ats_v1"},
    )
    assert ats_preview.status_code == 200
    ats_payload = ats_preview.json()
    assert any(
        line.startswith("- Rewritten (ats):")
        for section in ats_payload["sections"]
        for line in section["lines"]
    )

    stored_resume = next(doc for doc in db["tailored_resumes"].docs if str(doc["_id"]) == payload["id"])
    assert str(stored_resume["resume_snapshot_id"]) == snapshot_id
    assert stored_resume["template_source"] == "user_resume"
    assert stored_resume["template"] == "modern_v1"

    rewrite = client.post(f"/tailor/{payload['id']}/rewrite", headers=headers, json={"focus": "ats"})
    assert rewrite.status_code == 200
    assert rewrite.json()["rewritten_count"] > 0


def test_tailored_resume_can_reword_uploaded_resume_without_changing_section_layout(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    db = test_context["db"]

    resume_text = "\n".join(
        [
            "Taylor Candidate",
            "SUMMARY",
            "Product-minded engineer with strong analytics communication.",
            "SKILLS",
            "Python, SQL, Dashboards",
            "EXPERIENCE",
            "Analytics Engineer Intern",
            "- Built internal dashboards for operations teams.",
            "- Delivered API endpoints for reporting workflows.",
            "EDUCATION",
            "BS Computer Science",
        ]
    )
    ingest = client.post("/ingest/resume/text", headers=headers, json={"user_id": user_id, "text": resume_text})
    assert ingest.status_code == 200
    snapshot_id = ingest.json()["snapshot_id"]

    job = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "Analytics Platform Engineer",
            "company": "SkillBridge",
            "location": "Remote",
            "text": "Need strong dashboard delivery, analytics APIs, stakeholder communication, and Python experience." * 2,
        },
    )
    assert job.status_code == 200

    preview = client.post(
        "/tailor/preview",
        headers=headers,
        json={
            "user_id": user_id,
            "job_id": job.json()["id"],
            "resume_snapshot_id": snapshot_id,
            "template": "uploaded_resume_reword_v1",
        },
    )
    assert preview.status_code == 200
    payload = preview.json()
    section_titles = [section["title"] for section in payload["sections"]]
    assert section_titles == ["Header", "Summary", "Skills", "Experience", "Education"]
    assert "Targeted Highlights" not in section_titles
    assert payload["plain_text"].startswith("Taylor Candidate\nSUMMARY\n")
    assert "HEADER\n" not in payload["plain_text"]
    experience_section = next(section for section in payload["sections"] if section["title"] == "Experience")
    assert any(line.startswith("- Rewritten (balanced):") for line in experience_section["lines"])
    assert payload["template"] == "uploaded_resume_reword_v1"

    stored_resume = next(doc for doc in db["tailored_resumes"].docs if str(doc["_id"]) == payload["id"])
    assert stored_resume["template"] == "uploaded_resume_reword_v1"
    assert stored_resume["plain_text"].startswith("Taylor Candidate\nSUMMARY\n")

    rewrite = client.post(f"/tailor/{payload['id']}/rewrite", headers=headers, json={"focus": "ats"})
    assert rewrite.status_code == 200
    assert "- Rewritten (ats):" in rewrite.json()["plain_text"]

    docx = client.get(f"/tailor/{payload['id']}/export/docx", headers=headers)
    assert docx.status_code == 200

    pdf = client.get(f"/tailor/{payload['id']}/export/pdf", headers=headers)
    assert pdf.status_code == 200


def test_tailored_resume_can_preserve_editable_docx_formatting(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]

    source_doc = Document()
    source_doc.add_paragraph("Taylor Candidate")
    source_doc.add_paragraph("SUMMARY")
    source_doc.add_paragraph("Product-minded engineer with strong analytics communication.")
    source_doc.add_paragraph("EXPERIENCE")
    source_doc.add_paragraph("Analytics Engineer Intern")
    source_doc.add_paragraph("Built internal dashboards for operations teams.", style="List Bullet")
    source_doc.add_paragraph("Delivered API endpoints for reporting workflows.", style="List Bullet")
    buffer = BytesIO()
    source_doc.save(buffer)
    buffer.seek(0)

    ingest = client.post(
        "/ingest/resume/docx",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"user_id": user_id},
    )
    assert ingest.status_code == 200
    snapshot_id = ingest.json()["snapshot_id"]

    job = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "Analytics Platform Engineer",
            "company": "SkillBridge",
            "location": "Remote",
            "text": "Need strong dashboard delivery, analytics APIs, stakeholder communication, and Python experience." * 2,
        },
    )
    assert job.status_code == 200

    preview = client.post(
        "/tailor/preview",
        headers=headers,
        json={
            "user_id": user_id,
            "job_id": job.json()["id"],
            "resume_snapshot_id": snapshot_id,
            "template": "uploaded_resume_reword_v1",
        },
    )
    assert preview.status_code == 200
    tailored_id = preview.json()["id"]

    rewrite = client.post(f"/tailor/{tailored_id}/rewrite", headers=headers, json={"focus": "ats"})
    assert rewrite.status_code == 200

    exported = client.get(f"/tailor/{tailored_id}/export/docx", headers=headers)
    assert exported.status_code == 200

    exported_doc = Document(BytesIO(exported.content))
    list_paragraphs = [paragraph for paragraph in exported_doc.paragraphs if str(getattr(paragraph.style, "name", "")).startswith("List")]
    assert len(list_paragraphs) >= 2
    assert any("Rewritten (ats):" in paragraph.text for paragraph in list_paragraphs)


def test_job_match_required_coverage_and_missing_skills_are_populated(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_python = test_context["skill_python"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [{"skill_id": skill_python, "proficiency": 4}],
            "rejected": [],
            "edited": [],
        },
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "\n".join(
                [
                    "Required Qualifications:",
                    "- Python",
                    "- Machine Learning",
                    "Responsibilities:",
                    "- Build analytics dashboards and deliver APIs",
                ]
            ),
        },
    )
    assert ingest.status_code == 200

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": ingest.json()["id"]})
    assert match.status_code == 200
    payload = match.json()

    assert payload["required_skill_count"] == 2
    assert payload["required_matched_count"] == 1
    assert payload["matched_skill_count"] == 1
    assert payload["missing_skill_count"] == 1
    assert "Python" in payload["matched_skills"]
    assert "ML" in payload["missing_skills"]


def test_saved_job_history_does_not_persist_missing_skills(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_python = test_context["skill_python"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [{"skill_id": skill_python, "proficiency": 4}],
            "rejected": [],
            "edited": [],
        },
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Required skills: Python and Machine Learning. Responsibilities include analytics dashboards and API delivery.",
        },
    )
    assert ingest.status_code == 200

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": ingest.json()["id"]})
    assert match.status_code == 200
    live_payload = match.json()
    history_id = live_payload["history_id"]

    assert live_payload["missing_skill_count"] == 1
    assert "ML" in live_payload["missing_skills"]
    assert live_payload["gap_insights"]

    history = client.get("/tailor/history", headers=headers)
    assert history.status_code == 200
    history_entry = next(item for item in history.json() if item["id"] == history_id)
    assert history_entry["missing_skills"] == []

    detail = client.get(f"/tailor/history/{history_id}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["analysis"]["missing_skills"] == []
    assert detail.json()["analysis"]["missing_skill_count"] == 0
    assert detail.json()["analysis"]["gap_reasoning_summary"] == ""
    assert detail.json()["analysis"]["gap_insights"] == []


def test_job_match_falls_back_to_extracted_skills_when_required_section_is_not_detected(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_python = test_context["skill_python"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [{"skill_id": skill_python, "proficiency": 4}],
            "rejected": [],
            "edited": [],
        },
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "Analytics Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "We are looking for Python and Machine Learning experience to build analytics dashboards and APIs.",
        },
    )
    assert ingest.status_code == 200

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": ingest.json()["id"]})
    assert match.status_code == 200
    payload = match.json()

    assert payload["extracted_skill_count"] == 2
    assert payload["required_skill_count"] == 2
    assert payload["required_matched_count"] == 1
    assert payload["matched_skill_count"] == 1
    assert payload["missing_skill_count"] == 1


def test_reanalyze_recomputes_job_skills_from_saved_job_text(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_python = test_context["skill_python"]
    skill_ml = test_context["skill_ml"]
    db = test_context["db"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [
                {"skill_id": skill_python, "proficiency": 4},
                {"skill_id": skill_ml, "proficiency": 3},
            ],
            "rejected": [],
            "edited": [],
        },
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Required skills: Python and Machine Learning. Responsibilities include analytics dashboards and API delivery.",
        },
    )
    assert ingest.status_code == 200
    job_id = ingest.json()["id"]

    job_doc = next(doc for doc in db["job_ingests"].docs if str(doc["_id"]) == job_id)
    job_doc["extracted_skills"] = [{"skill_id": skill_ml, "skill_name": "ML", "matched_on": "alias", "count": 1}]
    job_doc["keywords"] = ["ml"]

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": job_id})
    assert match.status_code == 200
    payload = match.json()

    assert payload["extracted_skill_count"] >= 2
    assert "Python" in payload["matched_skills"]
    assert "ML" in payload["matched_skills"]


def test_reanalyze_dedupes_alias_equivalent_job_skills(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_ml = test_context["skill_ml"]
    db = test_context["db"]

    machine_learning_skill_id = "507f1f77bcf86cd799439099"
    db["skills"].docs.append(
        {
            "_id": ObjectId(machine_learning_skill_id),
            "name": "Machine Learning",
            "category": "Data",
            "categories": ["Data"],
            "aliases": ["ML"],
            "tags": ["ai"],
            "origin": "default",
            "hidden": False,
            "created_at": db["skills"].docs[0]["created_at"],
            "updated_at": db["skills"].docs[0]["updated_at"],
        }
    )

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [{"skill_id": skill_ml, "proficiency": 3}],
            "rejected": [],
            "edited": [],
        },
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Required skills: Machine Learning and ML. The role also includes model delivery, experimentation, and analytics communication.",
        },
    )
    assert ingest.status_code == 200

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": ingest.json()["id"]})
    assert match.status_code == 200
    payload = match.json()

    assert payload["matched_skill_count"] == 1
    assert payload["missing_skill_count"] == 0
    assert payload["required_skill_count"] == 1


def test_reanalyze_history_endpoint_returns_fresh_analysis(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    skill_python = test_context["skill_python"]
    skill_ml = test_context["skill_ml"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [
                {"skill_id": skill_python, "proficiency": 4},
                {"skill_id": skill_ml, "proficiency": 3},
            ],
            "rejected": [],
            "edited": [],
        },
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Required skills: Python and Machine Learning. Responsibilities include analytics dashboards and API delivery.",
        },
    )
    assert ingest.status_code == 200

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": ingest.json()["id"]})
    assert match.status_code == 200
    history_id = match.json()["history_id"]

    reanalyzed = client.post(f"/tailor/history/{history_id}/reanalyze", headers=headers, params={"user_id": user_id})
    assert reanalyzed.status_code == 200
    payload = reanalyzed.json()

    assert payload["history_id"]
    assert payload["history_id"] != history_id
    assert payload["matched_skill_count"] >= 1


def test_history_payload_drops_deleted_skill_names(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    db = test_context["db"]

    db["skills"].docs = [doc for doc in db["skills"].docs if doc.get("name") != "ML"]
    deleted_skill_id = ObjectId()
    db["job_match_runs"].docs.append(
        {
            "_id": ObjectId(),
            "user_id": ObjectId(test_context["user_id"]),
            "job_id": ObjectId(),
            "title": "Old Match",
            "company": "Acme",
            "location": "Remote",
            "text_preview": "Old saved analysis",
            "job_text_snapshot": "Machine learning and analytics work.",
            "analysis": {
                "job_id": "stale-job-id",
                "match_score": 62.0,
                "match_confidence_label": "Developing",
                "analysis_summary": "Saved summary",
                "resume_snapshot_id": None,
                "template_source": "default_template",
                "ignored_skill_names": [],
                "added_from_missing_skills": [],
                "matched_skill_ids": [str(deleted_skill_id)],
                "matched_skills": ["ML"],
                "missing_skill_ids": [str(deleted_skill_id)],
                "missing_skills": ["ML"],
                "strength_areas": [],
                "related_skills": ["ML"],
                "gap_insights": [
                    {
                        "skill_id": str(deleted_skill_id),
                        "skill_name": "ML",
                        "gap_type": "required",
                        "severity": "high",
                        "reason": "Old deleted skill",
                        "recommended_action": "Ignore this deleted skill",
                    }
                ],
                "score_breakdown": [],
                "recommended_next_steps": [],
                "extracted_skill_count": 1,
                "confirmed_skill_count": 0,
                "required_skill_count": 1,
                "required_matched_count": 0,
                "preferred_skill_count": 0,
                "preferred_matched_count": 0,
                "evidence_aligned_count": 0,
                "evidence_gap_count": 1,
                "keyword_overlap_count": 0,
                "keyword_overlap_terms": [],
                "semantic_alignment_score": 0,
                "semantic_alignment_explanation": "",
                "personal_skill_vector_score": 0,
                "personal_skill_vector_explanation": "",
                "retrieved_context": [],
            },
            "created_at": db["skills"].docs[0]["created_at"],
            "updated_at": db["skills"].docs[0]["updated_at"],
        }
    )

    history = client.get("/tailor/history", headers=headers)
    assert history.status_code == 200
    entry = next(item for item in history.json() if item["title"] == "Old Match")
    assert entry["matched_skills"] == []
    assert entry["missing_skills"] == []
    assert entry["related_skills"] == []

    detail = client.get(f"/tailor/history/{entry['id']}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["analysis"]["matched_skills"] == []
    assert detail.json()["analysis"]["missing_skills"] == []
    assert detail.json()["analysis"]["related_skills"] == []
    assert detail.json()["analysis"]["gap_insights"] == []


def test_tailored_resume_can_use_resume_evidence_as_template_source(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    db = test_context["db"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [
                {"skill_id": test_context["skill_python"], "proficiency": 4},
                {"skill_id": test_context["skill_ml"], "proficiency": 3},
            ],
            "rejected": [],
            "edited": [],
        },
    )

    resume_evidence = client.post(
        "/evidence/",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "resume",
            "title": "resume_final_v7.docx",
            "source": "manual-entry",
            "text_excerpt": "\n".join(
                [
                    "Jane Student",
                    "SUMMARY",
                    "Python and ML builder focused on analytics systems.",
                    "EXPERIENCE",
                    "Capstone Research Assistant",
                    "- Built analytics dashboards in Python.",
                    "- Delivered model evaluation workflows.",
                    "EDUCATION",
                    "BS Computer Science",
                ]
            ),
            "skill_ids": [test_context["skill_python"], test_context["skill_ml"]],
            "origin": "user",
        },
    )
    assert resume_evidence.status_code == 200
    evidence_id = resume_evidence.json()["id"]

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Looking for Python and Machine Learning experience to build analytics dashboards and model delivery pipelines.",
        },
    )
    assert ingest.status_code == 200

    preview = client.post(
        "/tailor/preview",
        headers=headers,
        json={
            "user_id": user_id,
            "job_id": ingest.json()["id"],
            "resume_evidence_id": evidence_id,
            "template": "uploaded_resume_reword_v1",
        },
    )
    assert preview.status_code == 200
    payload = preview.json()

    assert payload["resume_evidence_id"] == evidence_id
    assert payload["resume_snapshot_id"] is None
    assert payload["template_source"] == "evidence_resume"
    assert payload["plain_text"]
    assert "resume_final_v7.docx" not in payload["plain_text"]

    stored_resume = next(doc for doc in db["tailored_resumes"].docs if str(doc["_id"]) == payload["id"])
    assert str(stored_resume["resume_evidence_id"]) == evidence_id
    assert stored_resume["template_source"] == "evidence_resume"
    assert stored_resume["template"] == "ats_v1"


def test_tailored_resume_rewords_file_like_evidence_titles_into_resume_headers(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    db = test_context["db"]

    client.post(
        "/skills/confirmations/",
        headers=headers,
        json={
            "resume_snapshot_id": None,
            "confirmed": [{"skill_id": test_context["skill_python"], "proficiency": 4}],
            "rejected": [],
            "edited": [],
        },
    )

    evidence = client.post(
        "/evidence/",
        headers=headers,
        json={
            "user_id": user_id,
            "type": "project",
            "title": "capstone_final_v4.pdf",
            "source": "manual-entry",
            "text_excerpt": "Built analytics dashboards in Python and delivered API reporting workflows for stakeholders.",
            "skill_ids": [test_context["skill_python"]],
            "origin": "user",
        },
    )
    assert evidence.status_code == 200

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "Analytics Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Need Python, analytics dashboards, stakeholder communication, and reporting APIs.",
        },
    )
    assert ingest.status_code == 200

    preview = client.post(
        "/tailor/preview",
        headers=headers,
        json={
            "user_id": user_id,
            "job_id": ingest.json()["id"],
            "template": "ats_v1",
        },
    )
    assert preview.status_code == 200
    payload = preview.json()

    assert "capstone_final_v4.pdf" not in payload["plain_text"]
    assert "Selected Project" in payload["plain_text"] or "Relevant Experience" in payload["plain_text"]

    stored_resume = next(doc for doc in db["tailored_resumes"].docs if str(doc["_id"]) == payload["id"])
    assert "capstone_final_v4.pdf" not in stored_resume["plain_text"]


def test_job_match_ignores_short_aliases_for_long_skills(test_context):
    client = test_context["client"]
    headers = test_context["headers"]
    user_id = test_context["user_id"]
    db = test_context["db"]

    db["skills"].docs.append(
        {
            "_id": ObjectId(),
            "name": "Marketing Leadership",
            "category": "Business",
            "aliases": ["ML"],
            "tags": [],
            "origin": "user",
            "created_by_user_id": ObjectId(user_id),
            "hidden": False,
            "created_at": db["skills"].docs[0]["created_at"],
            "updated_at": db["skills"].docs[0]["updated_at"],
        }
    )

    ingest = client.post(
        "/tailor/job/ingest",
        headers=headers,
        json={
            "user_id": user_id,
            "title": "ML Engineer",
            "company": "Acme",
            "location": "Remote",
            "text": "Looking for ML engineers to build analytics systems and model delivery workflows.",
        },
    )
    assert ingest.status_code == 200

    match = client.post("/tailor/match", headers=headers, json={"user_id": user_id, "job_id": ingest.json()["id"]})
    assert match.status_code == 200
    payload = match.json()

    assert "Marketing Leadership" not in payload["missing_skills"]
    assert "Marketing Leadership" not in payload["matched_skills"]
