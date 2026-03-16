"""Job match, tailoring, retrieval, and export routes that turn user data into saved analyses and targeted resume outputs."""

from __future__ import annotations

import re
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from bson import ObjectId
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse

from app.core.auth import require_user
from app.core.config import settings
from app.core.db import get_db
from app.models.tailor import (
    AIPreferencesOut,
    AIPreferencesPatchIn,
    AISettingsDetailOut,
    AISettingsStatusOut,
    ExtractedSkill,
    GapInsight,
    JobIngestIn,
    JobIngestOut,
    JobMatchCompareOut,
    JobMatchHistoryDetailOut,
    JobMatchHistoryEntryOut,
    JobMatchOut,
    MatchScoreBreakdown,
    RAGContextItem,
    ResumeSection,
    RewriteBulletsIn,
    RewriteBulletsOut,
    TailoredResumeDetailOut,
    TailoredResumeListEntryOut,
    TailoredResumeOut,
    TailorPreviewIn,
    UserSkillVectorOut,
    UserSkillVectorHistoryPoint,
)
from app.utils.rag import retrieve_rag_context
from app.utils.skill_catalog import merge_skill_docs, normalize_skill_text
from app.utils.ai import (
    AVAILABLE_INFERENCE_MODES,
    cosine_similarity,
    embed_texts,
    get_inference_status,
    normalize_ai_preferences,
    rewrite_resume_bullets,
)
from app.utils.mongo import oid_str, ref_query, ref_values, to_object_id

router = APIRouter()
DEFAULT_RESUME_TEMPLATE_TEX = Path(__file__).resolve().parents[2] / "data" / "raw" / "default_resume_template.tex"

def now_utc():
    return datetime.now(timezone.utc)


def _scoped_user_id(user: dict, requested_user_id: str | None = None) -> str:
    effective_user_id = oid_str(user["_id"])
    candidate = str(requested_user_id or "").strip()
    if candidate and candidate != effective_user_id:
        raise HTTPException(status_code=403, detail="Forbidden for requested user_id")
    return effective_user_id


def _serialize_ai_preferences(user_doc: dict | None) -> AIPreferencesOut:
    prefs = normalize_ai_preferences((user_doc or {}).get("ai_preferences"))
    return AIPreferencesOut(
        inference_mode=prefs["inference_mode"],
        embedding_model=prefs["embedding_model"],
        zero_shot_model=prefs["zero_shot_model"],
        available_inference_modes=list(AVAILABLE_INFERENCE_MODES),
        available_embedding_models=settings.local_embedding_model_options_list,
        available_zero_shot_models=settings.local_zero_shot_model_options_list,
    )


def _build_ai_settings_detail(user_doc: dict | None) -> AISettingsDetailOut:
    prefs = _serialize_ai_preferences(user_doc)
    status = get_inference_status(
        {
            "inference_mode": prefs.inference_mode,
            "embedding_model": prefs.embedding_model,
            "zero_shot_model": prefs.zero_shot_model,
        }
    )
    return AISettingsDetailOut(
        preferences=prefs,
        **status,
    )


@router.get("/rag/search", response_model=list[RAGContextItem])
async def search_rag_context(q: str, user=Depends(require_user), limit: int = 5):
    # This endpoint is mainly for product/debug visibility: it lets the frontend show
    # which grounded snippets the retriever considers most relevant for a free-text query.
    db = get_db()
    preferences = normalize_ai_preferences((user or {}).get("ai_preferences"))
    rows = await retrieve_rag_context(
        db,
        user_id=oid_str(user["_id"]),
        query_text=q,
        preferences=preferences,
        limit=max(1, min(limit, 10)),
        source_types=("evidence", "resume_snapshot"),
    )
    return [RAGContextItem(**row) for row in rows]


@router.get("/user-vector", response_model=UserSkillVectorOut)
async def get_user_skill_vector(user=Depends(require_user)):
    db = get_db()
    preferences = normalize_ai_preferences((user or {}).get("ai_preferences"))
    doc = await _compute_and_store_user_skill_vector(db, oid_str(user["_id"]), ai_preferences=preferences)
    return UserSkillVectorOut(
        user_id=oid_str(user["_id"]),
        embedding_dimensions=int(doc.get("embedding_dimensions") or 0),
        confirmed_skill_count=int(doc.get("confirmed_skill_count") or 0),
        evidence_item_count=int(doc.get("evidence_item_count") or 0),
        portfolio_item_count=int(doc.get("portfolio_item_count") or 0),
        source_preview=str(doc.get("source_preview") or ""),
        updated_at=doc.get("updated_at"),
    )


@router.get("/user-vector/history", response_model=list[UserSkillVectorHistoryPoint])
async def get_user_skill_vector_history(user=Depends(require_user), limit: int = 12):
    db = get_db()
    rows = await db["user_skill_vector_history"].find(
        ref_query("user_id", oid_str(user["_id"])),
        {"score": 1, "label": 1, "updated_at": 1},
    ).sort("updated_at", -1).limit(max(1, min(limit, 30))).to_list(length=max(1, min(limit, 30)))
    ordered = list(reversed(rows))
    return [
        UserSkillVectorHistoryPoint(
            score=float(row.get("score") or 0.0),
            label=str(row.get("label") or ""),
            updated_at=row.get("updated_at"),
        )
        for row in ordered
    ]

def _job_focus_lines(text: str) -> list[tuple[str, float]]:
    section_markers = {
        "required": 3.5,
        "requirements": 3.5,
        "minimum qualifications": 3.5,
        "basic qualifications": 3.5,
        "must have": 3.5,
        "must-have": 3.5,
        "preferred": 2.5,
        "preferred qualifications": 2.5,
        "nice to have": 2.5,
        "nice-to-have": 2.5,
        "responsibilities": 2.25,
        "what you'll do": 2.25,
        "what you will do": 2.25,
        "in this role": 2.0,
        "experience with": 2.0,
    }

    active_weight = 1.0
    lines: list[tuple[str, float]] = []
    for raw_line in (text or "").splitlines():
        line = raw_line.strip(" \t-•*")
        if not line:
            continue
        lower = normalize_skill_text(line)
        matched_weight = next((weight for marker, weight in section_markers.items() if marker in lower), None)
        if matched_weight is not None:
            active_weight = matched_weight
            continue
        weight = active_weight
        if raw_line.lstrip().startswith(("-", "•", "*")):
            weight += 0.5
        if len(line) >= 18:
            lines.append((line, weight))
    return lines

def _normalize_keyword_phrase(value: str) -> str:
    value = normalize_skill_text(value)
    return value.strip(" ,;:/")

def _tokenize_keywords(text: str, extracted: Iterable[ExtractedSkill] | None = None) -> list[str]:
    generic_terms = {
        "experience", "skills", "skill", "requirements", "requirement", "responsibilities",
        "responsibility", "qualified", "qualification", "qualifications", "candidate",
        "candidates", "ability", "abilities", "knowledge", "strong", "excellent", "preferred",
        "required", "role", "team", "work", "working", "environment", "business", "company",
        "position", "opportunity", "including", "ability to", "experience in", "experience with",
    }
    stop_words = {
        "and", "the", "for", "with", "from", "that", "this", "have", "will", "your", "able",
        "must", "plus", "also", "using", "used", "into", "over", "such", "they", "their",
        "them", "than", "then", "only", "when", "where", "what", "were", "been", "being",
        "more", "less", "some", "many", "each", "make", "made", "through", "within", "across",
        "under", "while", "about", "other", "need", "needs", "build", "built",
    }

    phrase_scores: dict[str, float] = {}

    def add_phrase(raw_value: str, weight: float):
        phrase = _normalize_keyword_phrase(raw_value)
        if not phrase or phrase in generic_terms:
            return
        tokens = [token for token in phrase.split() if token not in stop_words]
        if not tokens:
            return
        if len(tokens) == 1 and len(tokens[0]) < 4 and "+" not in tokens[0] and "#" not in tokens[0]:
            return
        candidate = " ".join(tokens[:4])
        if candidate in generic_terms:
            return
        phrase_scores[candidate] = max(phrase_scores.get(candidate, 0.0), 0.0) + weight

    for skill in extracted or []:
        add_phrase(skill.skill_name, 7.0)

    for line, weight in _job_focus_lines(text):
        lower = line.lower()
        for segment in re.split(r"[;,/]|(?:\s+-\s+)", lower):
            segment = segment.strip()
            if not segment:
                continue
            add_phrase(segment, weight)
            words = re.findall(r"[A-Za-z][A-Za-z0-9+.#-]{2,}", segment)
            filtered = [word for word in words if word not in stop_words and word not in generic_terms]
            for word in filtered:
                add_phrase(word, weight * 0.65)
            for i in range(max(0, len(filtered) - 1)):
                add_phrase(" ".join(filtered[i : i + 2]), weight * 0.8)

    scored = sorted(phrase_scores.items(), key=lambda item: (-item[1], -len(item[0]), item[0]))
    return [phrase for phrase, _score in scored[:35]]

def _top_average(scores: Iterable[float], limit: int) -> float:
    positives = [score for score in scores if score > 0]
    if not positives:
        return 0.0
    ranked = sorted(positives, reverse=True)[:limit]
    return sum(ranked) / len(ranked)


async def _load_latest_resume_snapshot_text(db, user_id: str) -> str:
    snapshot = await db["resume_snapshots"].find_one(
        ref_query("user_id", user_id),
        sort=[("created_at", -1)],
    )
    return str((snapshot or {}).get("raw_text") or "").strip()


def _build_user_vector_source_text(
    confirmed_entries: list[dict],
    skill_docs: dict[str, dict],
    items: list[dict],
    resume_text: str,
) -> tuple[str, dict]:
    confirmed_skill_names = [
        _display_skill_name(skill_docs.get(str(entry.get("skill_id") or "").strip()), str(entry.get("skill_id") or "").strip())
        for entry in confirmed_entries
        if str(entry.get("skill_id") or "").strip()
    ]
    confirmed_skill_names = [name for name in _dedupe_preserve_order(confirmed_skill_names) if name]

    evidence_items = [item for item in items if str(item.get("_source_collection") or "") == "evidence"]

    sections: list[str] = []
    if confirmed_skill_names:
        sections.append(f"Confirmed skills: {', '.join(confirmed_skill_names[:40])}.")
    for item in items[:10]:
        blob = _normalize_text_blob(item)
        if blob:
            sections.append(blob[:500])
    if resume_text:
        sections.append(resume_text[:1500])

    source_text = "\n".join(section for section in sections if section).strip()
    return source_text, {
        "confirmed_skill_count": len(confirmed_skill_names),
        "evidence_item_count": len(evidence_items),
        "portfolio_item_count": 0,
    }


async def _compute_and_store_user_skill_vector(db, user_id: str, ai_preferences: dict | None = None) -> dict:
    conf = await _load_profile_confirmation(db, user_id)
    confirmed_entries = list((conf or {}).get("confirmed") or [])
    confirmed_skill_ids = [str(entry.get("skill_id") or "").strip() for entry in confirmed_entries if str(entry.get("skill_id") or "").strip()]
    skill_docs = await _load_skills_by_ids(db, confirmed_skill_ids)
    items = await _load_user_items(db, user_id)
    resume_text = await _load_latest_resume_snapshot_text(db, user_id)
    source_text, stats = _build_user_vector_source_text(confirmed_entries, skill_docs, items, resume_text)

    vectors, provider = await embed_texts([source_text or "empty profile"], preferences=ai_preferences)
    vector = vectors[0] if vectors else []
    doc = {
        "user_id": to_object_id(user_id),
        "provider": provider,
        "vector": vector,
        "embedding_dimensions": len(vector),
        "source_preview": source_text[:300],
        **stats,
        "updated_at": now_utc(),
    }
    await db["user_skill_vectors"].update_one(
        ref_query("user_id", user_id),
        {"$set": doc},
        upsert=True,
    )
    await db["user_skill_vector_history"].insert_one(
        {
            "user_id": to_object_id(user_id),
            "score": round(min(100.0, len(vector) * 5.0), 2),
            "label": f"{stats['confirmed_skill_count']} skills / {stats['evidence_item_count']} evidence",
            "updated_at": doc["updated_at"],
        }
    )
    return doc

def _is_hidden_skill_doc(doc: dict) -> bool:
    name = (doc.get("name") or "").strip().lower()
    if not name:
        return True
    if doc.get("hidden") is True:
        return True
    blocked_terms = ("test", "demo", "sample", "mock", "dummy", "placeholder")
    return any(term in name for term in blocked_terms)

async def _load_skill_catalog(db) -> list[dict]:
    cursor = db["skills"].find({}, {"name": 1, "aliases": 1, "category": 1, "hidden": 1})
    docs = await cursor.to_list(length=5000)
    return merge_skill_docs([doc for doc in docs if not _is_hidden_skill_doc(doc)])

def _match_skills(job_text: str, skills: Iterable[dict]) -> list[ExtractedSkill]:
    text = normalize_skill_text(job_text)
    matches: dict[str, ExtractedSkill] = {}

    def bump(sid: str, name: str, matched_on: str):
        key = sid
        if key not in matches:
            matches[key] = ExtractedSkill(skill_id=sid, skill_name=name, matched_on=matched_on, count=1)
        else:
            matches[key].count += 1

    for s in skills:
        sid = str(s["_id"])
        name = (s.get("name") or "").strip()
        if not name:
            continue

        # word-boundary match for names; allow special tokens like C++ / C# via loose matching
        n = normalize_skill_text(name)
        if len(n) >= 2:
            if re.search(rf"(?<![A-Za-z0-9]){re.escape(n)}(?![A-Za-z0-9])", text):
                bump(sid, name, "name")

        for a in s.get("aliases", []) or []:
            a = (a or "").strip()
            if not a:
                continue
            al = normalize_skill_text(a)
            if re.search(rf"(?<![A-Za-z0-9]){re.escape(al)}(?![A-Za-z0-9])", text):
                bump(sid, name, "alias")

    # sort by count desc then name
    return sorted(matches.values(), key=lambda x: (-x.count, x.skill_name.lower()))

async def _load_user_items(db, user_id: str) -> list[dict]:
    user_filter = ref_query("user_id", user_id)
    items: list[dict] = []
    evidence_docs = await db["evidence"].find(user_filter).sort("updated_at", -1).to_list(length=1000)
    for evidence in evidence_docs:
        items.append(
            {
                "_id": evidence["_id"],
                "type": f"evidence:{evidence.get('type', 'other')}",
                "title": evidence.get("title", "") or "Evidence",
                "org": evidence.get("org"),
                "summary": evidence.get("summary") or evidence.get("text_excerpt") or "",
                "bullets": evidence.get("bullets", []) or [],
                "links": evidence.get("links", []) or ([evidence.get("source")] if evidence.get("source") else []),
                "skill_ids": evidence.get("skill_ids", []) or [],
                "priority": evidence.get("priority", 1),
                "updated_at": evidence.get("updated_at"),
                "created_at": evidence.get("created_at"),
                "_source_collection": "evidence",
            }
        )
    return items

def _score_item(item: dict, job_skill_ids: set[str], keywords: set[str]) -> float:
    sids = {str(sid) for sid in (item.get("skill_ids", []) or [])}
    overlap = len(sids & job_skill_ids)
    pri = float(item.get("priority", 0) or 0)

    # keyword overlap from title + summary + bullets
    txt = normalize_skill_text(" ".join(
        [item.get("title",""), item.get("summary","")] + (item.get("bullets", []) or [])
    ))
    kw_hit = sum(1 for k in keywords if k in txt)

    return overlap * 5.0 + kw_hit * 1.0 + pri * 0.25


def _attach_retrieved_context_to_items(items: list[dict], retrieved_context: list[dict]) -> dict[str, float]:
    # Evidence items still drive resume selection, but retrieval gives us a better way
    # to boost items whose text is semantically close to the job posting. We attach the
    # winning snippets back onto those items so later summarization can reuse them.
    context_score_by_item_id: dict[str, float] = {}
    snippets_by_item_id: dict[str, list[str]] = {}
    for context in retrieved_context:
        source_type = str(context.get("source_type") or "").strip()
        if source_type != "evidence":
            continue
        source_id = str(context.get("source_id") or "").strip()
        if not source_id:
            continue
        snippets_by_item_id.setdefault(source_id, [])
        snippet = str(context.get("snippet") or "").strip()
        if snippet:
            snippets_by_item_id[source_id].append(snippet)
        context_score_by_item_id[source_id] = max(
            float(context_score_by_item_id.get(source_id, 0.0)),
            float(context.get("score") or 0.0),
        )

    for item in items:
        item_id = oid_str(item.get("_id"))
        if not item_id:
            continue
        if item_id in snippets_by_item_id:
            item["rag_snippets"] = _dedupe_preserve_order(snippets_by_item_id[item_id])[:3]
    return context_score_by_item_id

def _dedupe_preserve_order(values: Iterable[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = str(value or "").strip()
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(text)
    return out

def _normalize_text_blob(item: dict) -> str:
    parts = [item.get("title", ""), item.get("summary", ""), item.get("org", "")]
    parts.extend(item.get("bullets", []) or [])
    parts.extend(item.get("links", []) or [])
    return normalize_skill_text(" ".join(str(part or "") for part in parts))

async def _load_profile_confirmation(db, user_id: str) -> dict | None:
    confirmation = await db["resume_skill_confirmations"].find_one(
        {"user_id": {"$in": ref_values(user_id)}, "resume_snapshot_id": None},
        sort=[("created_at", -1)],
    )
    if confirmation:
        return confirmation
    return await db["resume_skill_confirmations"].find_one(
        {"user_id": {"$in": ref_values(user_id)}},
        sort=[("created_at", -1)],
    )

async def _load_skills_by_ids(db, skill_ids: Iterable[str]) -> dict[str, dict]:
    oids: list[ObjectId] = []
    raw_ids: list[str] = []
    for skill_id in skill_ids:
        text = str(skill_id or "").strip()
        if not text:
            continue
        raw_ids.append(text)
        oid = to_object_id(text) if ObjectId.is_valid(text) else None
        if oid is not None:
            oids.append(oid)
    if not raw_ids:
        return {}

    query_values: list[object] = list(oids)
    query_values.extend(text for text in raw_ids if text not in {str(oid) for oid in oids})
    docs = await db["skills"].find(
        {"_id": {"$in": query_values}},
        {"name": 1, "category": 1, "aliases": 1, "hidden": 1},
    ).to_list(length=max(len(query_values), 1))
    visible_docs = [doc for doc in docs if not _is_hidden_skill_doc(doc)]
    merged_docs = merge_skill_docs(visible_docs)
    out: dict[str, dict] = {}
    for doc in merged_docs:
        merged_ids = list(doc.get("merged_ids") or [])
        if not merged_ids and doc.get("_id") is not None:
            merged_ids = [oid_str(doc.get("_id"))]
        for merged_id in merged_ids:
            text = str(merged_id or "").strip()
            if text:
                out[text] = doc
    return out

def _skill_terms(doc: dict | None) -> set[str]:
    if not doc:
        return set()
    values = [(doc.get("name") or "").strip()]
    values.extend(str(alias or "").strip() for alias in (doc.get("aliases") or []))
    terms: set[str] = set()
    for value in values:
        if not value:
            continue
        normalized = normalize_skill_text(value)
        if normalized:
            terms.add(normalized)
    return terms

def _score_percent(numerator: int, denominator: int) -> float:
    if denominator <= 0:
        return 0.0
    return round((numerator / denominator) * 100.0, 2)

def _clamp_score(value: float) -> float:
    return round(max(0.0, min(100.0, value)), 2)

def _skill_match_pattern(value: str) -> str:
    return rf"(?<![A-Za-z0-9]){re.escape(normalize_skill_text(value))}(?![A-Za-z0-9])"


def _normalize_ignored_skill_names(values: Iterable[str] | None) -> list[str]:
    return _dedupe_preserve_order(str(value or "").strip() for value in (values or []) if str(value or "").strip())


def _is_ignored_skill(
    skill_id: str,
    ignored_terms: set[str],
    skill_docs: dict[str, dict],
    extracted_name: str = "",
) -> bool:
    if not ignored_terms:
        return False
    candidates = set()
    if extracted_name:
        normalized = normalize_skill_text(extracted_name)
        if normalized:
            candidates.add(normalized)
    display_name = normalize_skill_text(_display_skill_name(skill_docs.get(skill_id), skill_id))
    if display_name:
        candidates.add(display_name)
    candidates.update(_skill_terms(skill_docs.get(skill_id)))
    return bool(candidates & ignored_terms)

def _classify_job_skill_priority(job_text: str, extracted: list[dict], skill_docs: dict[str, dict]) -> tuple[set[str], set[str]]:
    required_markers = (
        "required", "requirements", "must have", "must-have", "minimum qualifications",
        "basic qualifications", "you will need", "need to have"
    )
    preferred_markers = (
        "preferred", "nice to have", "nice-to-have", "bonus", "preferred qualifications",
        "ideal", "plus", "would be a plus"
    )

    current_section: str | None = None
    required_ids: set[str] = set()
    preferred_ids: set[str] = set()
    section_reset_markers = (
        "responsibilities", "what you'll do", "what you will do", "about the role",
        "benefits", "compensation", "about us", "why join", "who we are",
    )
    section_blocks: dict[str, list[str]] = {"required": [], "preferred": []}
    extracted_terms_by_id: dict[str, set[str]] = {}

    for entry in extracted:
        skill_id = str(entry.get("skill_id") or "")
        if not skill_id:
            continue
        terms = _skill_terms(skill_docs.get(skill_id))
        if terms:
            extracted_terms_by_id[skill_id] = terms

    for raw_line in (job_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            current_section = None
            continue
        lower = normalize_skill_text(line)
        if any(marker in lower for marker in required_markers):
            current_section = "required"
        elif any(marker in lower for marker in preferred_markers):
            current_section = "preferred"
        elif any(marker in lower for marker in section_reset_markers):
            current_section = None
        elif len(lower) < 60 and lower.endswith(":"):
            current_section = None

        if current_section in {"required", "preferred"}:
            section_blocks[current_section].append(lower)

        for skill_id, terms in extracted_terms_by_id.items():
            if not any(term and re.search(_skill_match_pattern(term), lower) for term in terms):
                continue
            if current_section == "required":
                required_ids.add(skill_id)
            elif current_section == "preferred":
                preferred_ids.add(skill_id)

    for section_name, target_ids in (("required", required_ids), ("preferred", preferred_ids)):
        block_text = "\n".join(section_blocks.get(section_name) or [])
        if not block_text:
            continue
        for skill_id, terms in extracted_terms_by_id.items():
            if any(term and re.search(_skill_match_pattern(term), block_text) for term in terms):
                target_ids.add(skill_id)

    if not required_ids:
        strong_required_blob = "\n".join(
            normalize_skill_text(raw_line.strip())
            for raw_line in (job_text or "").splitlines()
            if raw_line.strip()
            and (
                any(marker in normalize_skill_text(raw_line) for marker in required_markers)
                or re.search(r"\b(must|required|minimum)\b", normalize_skill_text(raw_line))
            )
        )
        if strong_required_blob:
            for skill_id, terms in extracted_terms_by_id.items():
                if any(term and re.search(_skill_match_pattern(term), strong_required_blob) for term in terms):
                    required_ids.add(skill_id)

    preferred_ids -= required_ids
    return required_ids, preferred_ids

def _match_confidence_label(score: float) -> str:
    if score >= 85:
        return "Strong"
    if score >= 70:
        return "Promising"
    if score >= 50:
        return "Moderate"
    return "Early"

def _build_strength_areas(matched_skill_docs: list[dict], evidence_backed_ids: set[str], item_hits: list[dict]) -> list[str]:
    category_counts: dict[str, int] = {}
    evidence_category_counts: dict[str, int] = {}
    for doc in matched_skill_docs:
        category = (doc.get("category") or "General").strip() or "General"
        category_counts[category] = category_counts.get(category, 0) + 1
        if oid_str(doc.get("_id")) in evidence_backed_ids:
            evidence_category_counts[category] = evidence_category_counts.get(category, 0) + 1

    ordered_categories = sorted(category_counts.items(), key=lambda item: (-item[1], item[0].lower()))
    strengths: list[str] = []
    for category, count in ordered_categories[:3]:
        evidence_count = evidence_category_counts.get(category, 0)
        if evidence_count > 0:
            strengths.append(f"{category} ({count} matched, {evidence_count} evidence-backed)")
        else:
            strengths.append(f"{category} ({count} matched skills)")

    if item_hits:
        top_item = item_hits[0]
        title = (top_item.get("title") or "").strip()
        item_type = str(top_item.get("type") or "work").replace("evidence:", "").replace("_", " ")
        if title:
            strengths.append(f"{item_type.title()} evidence from {title}")

    return strengths[:4]


def _build_gap_insights(
    missing_ids: list[str],
    required_ids: set[str],
    preferred_ids: set[str],
    related_skill_ids: list[str],
    skill_docs: dict[str, dict],
) -> tuple[str, list[GapInsight]]:
    related_set = set(related_skill_ids)
    insights: list[GapInsight] = []
    required_missing = 0
    bridgeable_missing = 0

    for skill_id in missing_ids:
        name = _display_skill_name(skill_docs.get(skill_id), skill_id)
        if skill_id in required_ids:
            gap_type = "required"
            severity = "high"
            reason = f"{name} appears in a likely required section of the posting and is not yet confirmed on the profile."
            action = f"Add direct evidence or confirm recent work that proves {name}."
            required_missing += 1
        elif skill_id in preferred_ids:
            gap_type = "preferred"
            severity = "medium"
            reason = f"{name} appears in a preferred or secondary section, so it improves competitiveness but is less critical than the required set."
            action = f"Strengthen the profile with a project, course, or evidence snippet tied to {name}."
        elif skill_id in related_set:
            gap_type = "adjacent"
            severity = "medium"
            reason = f"{name} is not confirmed directly, but the profile contains semantically related experience that could help bridge the gap."
            action = f"Translate related work into explicit {name} evidence or add the closest supporting project."
            bridgeable_missing += 1
        else:
            gap_type = "general"
            severity = "medium"
            reason = f"{name} was extracted from the posting but does not yet appear in confirmed skills or supporting evidence."
            action = f"Decide whether to learn {name} or remove it from this analysis if the extraction is not relevant."

        insights.append(
            GapInsight(
                skill_id=skill_id,
                skill_name=name,
                gap_type=gap_type,
                severity=severity,
                reason=reason,
                recommended_action=action,
            )
        )

    summary_parts: list[str] = []
    if required_missing:
        summary_parts.append(f"{required_missing} missing skills appear to be required, which is the main reason the score is being capped.")
    if bridgeable_missing:
        summary_parts.append(f"{bridgeable_missing} gaps are partially offset by adjacent confirmed experience, so they are more about proof and positioning than total absence.")
    remaining = max(0, len(missing_ids) - required_missing - bridgeable_missing)
    if remaining:
        summary_parts.append(f"{remaining} remaining gaps are lower-priority or secondary skills relative to the required set.")
    if not summary_parts:
        summary_parts.append("No material skill gaps were identified beyond the confirmed coverage already shown in the score.")

    return " ".join(summary_parts), insights[:8]

def _display_skill_name(doc: dict | None, fallback: str = "") -> str:
    if not doc:
        return fallback
    return str(doc.get("name") or fallback or "").strip()

_RESUME_SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "Header": (),
    "Summary": ("summary", "professional summary", "profile", "objective", "about"),
    "Skills": ("skills", "technical skills", "core skills", "competencies", "technical competencies"),
    "Experience": (
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "relevant experience",
    ),
    "Projects": ("projects", "relevant projects", "selected projects", "academic projects"),
    "Education": ("education", "academic background"),
    "Certifications": ("certifications", "licenses", "certificates"),
    "Leadership": ("leadership", "leadership experience", "activities", "campus involvement"),
    "Awards": ("awards", "honors", "honours"),
}


def _canonical_resume_section_title(line: str) -> str | None:
    normalized = normalize_skill_text(re.sub(r"[:\-|]+$", "", line or "").strip())
    if not normalized:
        return None
    for title, aliases in _RESUME_SECTION_ALIASES.items():
        if normalized == normalize_skill_text(title):
            return title
        if any(normalized == normalize_skill_text(alias) for alias in aliases):
            return title
    return None


def _looks_like_resume_heading(line: str) -> str | None:
    stripped = (line or "").strip()
    if not stripped or len(stripped) > 48:
        return None
    canonical = _canonical_resume_section_title(stripped)
    if canonical:
        return canonical
    plain = re.sub(r"[^A-Za-z/& ]", "", stripped).strip()
    if plain and plain.isupper():
        canonical = _canonical_resume_section_title(plain.title())
        if canonical:
            return canonical
    return None


async def _load_resume_template_snapshot(db, user_id: str, resume_snapshot_id: str | None) -> dict | None:
    if resume_snapshot_id:
        query = ref_query("_id", resume_snapshot_id)
        query.update(ref_query("user_id", user_id))
        return await db["resume_snapshots"].find_one(query)
    return await db["resume_snapshots"].find_one(
        ref_query("user_id", user_id),
        sort=[("created_at", -1)],
    )


def _parse_resume_sections(raw_text: str) -> list[ResumeSection]:
    sections: list[ResumeSection] = []
    current_title = "Header"
    current_lines: list[str] = []

    def flush():
        nonlocal current_lines, current_title
        cleaned = [line.strip() for line in current_lines if str(line or "").strip()]
        if cleaned:
            sections.append(ResumeSection(title=current_title, lines=cleaned))
        current_lines = []

    for raw_line in (raw_text or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            if current_lines and current_lines[-1] != "":
                current_lines.append("")
            continue
        heading = _looks_like_resume_heading(line)
        if heading:
            flush()
            current_title = heading
            continue
        current_lines.append(line)

    flush()
    return sections


def _build_default_header_lines(user_doc: dict | None) -> list[str]:
    username = str((user_doc or {}).get("username") or "").strip()
    email = str((user_doc or {}).get("email") or "").strip()
    display_name = username or (email.split("@")[0] if email else "SkillBridge User")
    header_lines = [display_name]
    if email:
        header_lines.append(email)
    return header_lines


def _looks_like_name_line(value: str) -> bool:
    text = _clean_resume_line(value)
    if not text or len(text) > 60:
        return False
    if "@" in text or "http" in text or re.search(r"\d{3}", text):
        return False
    tokens = [token for token in re.split(r"\s+", text) if token]
    if not 2 <= len(tokens) <= 4:
        return False
    alpha_tokens = [token for token in tokens if re.search(r"[A-Za-z]", token)]
    if len(alpha_tokens) != len(tokens):
        return False
    return all(token[:1].isupper() for token in alpha_tokens if token)


def _extract_resume_header_lines(raw_text: str, user_doc: dict | None = None) -> list[str]:
    fallback = _build_default_header_lines(user_doc)
    text = str(raw_text or "")
    if not text.strip():
        return fallback

    lines = [_clean_resume_line(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    header_window = lines[:12]

    email_match = re.search(r"\b[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}\b", text, re.IGNORECASE)
    phone_match = re.search(r"(?:\+?1[\s.\-]?)?(?:\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4})", text)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s|,;]+", text, re.IGNORECASE)
    github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/[^\s|,;]+", text, re.IGNORECASE)

    name = next((line for line in header_window if _looks_like_name_line(line)), "") or str((user_doc or {}).get("username") or "").strip()
    if not name and fallback:
        name = fallback[0]

    contact_parts: list[str] = []
    if phone_match:
        contact_parts.append(_clean_resume_line(phone_match.group(0)))
    email = _clean_resume_line(email_match.group(0)) if email_match else str((user_doc or {}).get("email") or "").strip()
    if email:
        contact_parts.append(email)
    profile_links: list[str] = []
    if linkedin_match:
        profile_links.append(_clean_resume_line(linkedin_match.group(0).rstrip(").,;")))
    if github_match:
        profile_links.append(_clean_resume_line(github_match.group(0).rstrip(").,;")))
    contact_parts.extend(profile_links)

    location_line = next(
        (
            line for line in header_window
            if line != name
            and line not in contact_parts
            and "@" not in line
            and not re.search(r"(linkedin|github|http)", line, re.IGNORECASE)
            and len(line) <= 48
            and re.search(r"[A-Za-z]", line)
        ),
        "",
    )

    header_lines: list[str] = []
    if name:
        header_lines.append(name)
    if location_line:
        header_lines.append(location_line)
    if contact_parts:
        header_lines.append(" | ".join(_dedupe_preserve_order(contact_parts)))

    return header_lines or fallback


def _canonical_template_section_title(raw_title: str) -> str:
    cleaned = _clean_resume_line(raw_title)
    canonical = _canonical_resume_section_title(cleaned)
    if canonical:
        return canonical
    normalized = normalize_skill_text(cleaned)
    if "projects" in normalized or "research" in normalized:
        return "Projects"
    if "experience" in normalized:
        return "Experience"
    if "skills" in normalized:
        return "Skills"
    if "education" in normalized:
        return "Education"
    if "cert" in normalized:
        return "Certifications"
    if "summary" in normalized or "profile" in normalized:
        return "Summary"
    return cleaned or "Section"


def _load_default_resume_template_sections(user_doc: dict | None, include_content: bool = True) -> list[ResumeSection]:
    if not DEFAULT_RESUME_TEMPLATE_TEX.exists():
        return []
    raw_tex = DEFAULT_RESUME_TEMPLATE_TEX.read_text(encoding="utf-8", errors="ignore")
    body_match = re.search(r"\\begin\{document\}(.*)\\end\{document\}", raw_tex, re.DOTALL)
    body = body_match.group(1) if body_match else raw_tex

    sections: list[ResumeSection] = []
    sections.append(ResumeSection(title="Header", lines=_build_default_header_lines(user_doc)))

    matches = list(re.finditer(r"\\section\*\{([^}]+)\}", body))
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(body)
        title = _canonical_template_section_title(match.group(1))
        lines = _latex_block_to_lines(body[start:end]) if include_content else []
        sections.append(ResumeSection(title=title, lines=lines))

    if include_content:
        return [section for section in sections if section.lines]
    return sections


def _clean_resume_line(line: str) -> str:
    return re.sub(r"\s+", " ", str(line or "").strip())


def _latex_block_to_lines(block: str) -> list[str]:
    text = str(block or "")
    if not text.strip():
        return []

    text = re.sub(r"%.*", "", text)
    text = re.sub(r"\$[^$]*\$", " ", text)
    text = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\hfill", " | ", text)
    text = re.sub(r"\\item\s*", "\n- ", text)
    text = re.sub(r"\\begin\{[^}]+\}", "\n", text)
    text = re.sub(r"\\end\{[^}]+\}", "\n", text)
    text = re.sub(r"\\\\(?:\[[^\]]*\])?", "\n", text)
    text = re.sub(r"\\[A-Za-z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    text = re.sub(r"\n{2,}", "\n", text)

    lines: list[str] = []
    for raw_line in text.splitlines():
        cleaned = _clean_resume_line(raw_line.strip(" |-"))
        if cleaned:
            lines.append(cleaned)
    return lines


def _looks_like_formula_noise(text: str) -> bool:
    value = str(text or "").strip()
    if not value:
        return True
    if re.search(r"\\(frac|sum|int|sqrt|alpha|beta|gamma|theta|lambda|sigma|mu|pi|cdot|times)", value):
        return True
    if "$" in value or re.search(r"[=<>^_]{2,}", value):
        return True
    non_word_ratio = sum(1 for char in value if not (char.isalnum() or char.isspace())) / max(1, len(value))
    alpha_ratio = sum(1 for char in value if char.isalpha()) / max(1, len(value))
    return non_word_ratio > 0.28 or alpha_ratio < 0.45


def _clean_evidence_text_for_resume(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"\$[^$]*\$", " ", value)
    value = re.sub(r"\\\((.*?)\\\)", " ", value)
    value = re.sub(r"\\\[(.*?)\\\]", " ", value)
    value = re.sub(r"\\(frac|sum|int|sqrt|alpha|beta|gamma|theta|lambda|sigma|mu|pi|cdot|times)\b", " ", value)
    value = re.sub(r"https?://\S+", " ", value)
    value = re.sub(r"\b[a-zA-Z]+_[a-zA-Z0-9_]+\b", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def _sentence_candidates(text: str) -> list[str]:
    cleaned = _clean_evidence_text_for_resume(text)
    if not cleaned:
        return []
    raw_parts = re.split(r"(?<=[.!?])\s+|;|\n|•", cleaned)
    candidates: list[str] = []
    for raw_part in raw_parts:
        line = _clean_resume_line(raw_part.strip(" -•"))
        word_count = len(line.split())
        if word_count < 6 or word_count > 32:
            continue
        if _looks_like_formula_noise(line):
            continue
        candidates.append(line.rstrip(".") + ".")
    return candidates


def _summarize_item_for_resume(item: dict, selected_skill_names: list[str], max_bullets_per_item: int) -> list[str]:
    text_sources = [
        *[str(snippet or "") for snippet in (item.get("rag_snippets") or [])],
        *[str(bullet or "") for bullet in (item.get("bullets") or [])],
        str(item.get("summary") or ""),
        str(item.get("title") or ""),
    ]
    candidates: list[tuple[float, str]] = []
    normalized_skills = [normalize_skill_text(name) for name in selected_skill_names if normalize_skill_text(name)]
    action_verbs = ("built", "implemented", "developed", "designed", "led", "improved", "analyzed", "created", "optimized", "deployed")

    for text in text_sources:
        for sentence in _sentence_candidates(text):
            lowered = normalize_skill_text(sentence)
            skill_hits = sum(1 for skill in normalized_skills if skill and skill in lowered)
            verb_hits = sum(1 for verb in action_verbs if verb in lowered)
            digit_bonus = 0.3 if re.search(r"\b\d+[%x]?\b", sentence) else 0.0
            score = (skill_hits * 2.0) + (verb_hits * 1.2) + digit_bonus - (len(sentence.split()) / 100.0)
            candidates.append((score, sentence))

    ordered: list[str] = []
    seen: set[str] = set()
    for _score, sentence in sorted(candidates, key=lambda item: item[0], reverse=True):
        key = normalize_skill_text(sentence)
        if not key or key in seen:
            continue
        seen.add(key)
        ordered.append(sentence)
        if len(ordered) >= max_bullets_per_item:
            break

    if ordered:
        return ordered

    fallback = _clean_resume_line(str(item.get("summary") or item.get("title") or ""))
    if fallback:
        if _looks_like_formula_noise(fallback):
            fallback = _clean_resume_line(str(item.get("title") or "Relevant work"))
        if fallback:
            return [fallback.rstrip(".") + "."]
    return []


def _is_project_like_item(item: dict) -> bool:
    item_type = normalize_skill_text(str(item.get("type") or ""))
    title = normalize_skill_text(str(item.get("title") or ""))
    summary = normalize_skill_text(str(item.get("summary") or ""))
    return any(token in item_type for token in ("project", "portfolio", "research")) or any(
        token in title or token in summary for token in ("project", "prototype", "research", "capstone")
    )


def _chunk_skill_lines(skill_names: list[str], width: int = 92) -> list[str]:
    if not skill_names:
        return []
    lines: list[str] = []
    current = ""
    for name in skill_names:
        part = name.strip()
        if not part:
            continue
        candidate = part if not current else f"{current}, {part}"
        if current and len(candidate) > width:
            lines.append(current)
            current = part
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def _extract_original_skills(section: ResumeSection | None) -> list[str]:
    if not section:
        return []
    raw = ", ".join(line for line in section.lines if line.strip())
    parts = re.split(r"[,\n|/•]+", raw)
    out: list[str] = []
    seen: set[str] = set()
    for part in parts:
        cleaned = _clean_resume_line(part)
        key = normalize_skill_text(cleaned)
        if not cleaned or not key or key in seen:
            continue
        seen.add(key)
        out.append(cleaned)
    return out


def _build_targeted_summary_lines(
    base_section: ResumeSection | None,
    target_label: str,
    skill_names: list[str],
    selected_items: list[dict],
) -> list[str]:
    lines: list[str] = []
    seen: set[str] = set()
    for line in (base_section.lines if base_section else []):
        cleaned = _clean_resume_line(line)
        key = normalize_skill_text(cleaned)
        if cleaned and key not in seen:
            seen.add(key)
            lines.append(cleaned)
        if len(lines) >= 2:
            break

    highlight_skills = ", ".join(skill_names[:6])
    top_titles = ", ".join(_clean_resume_line(item.get("title", "")) for item in selected_items[:2] if item.get("title"))
    tailored_line_parts = [f"Tailored for {target_label}"]
    if highlight_skills:
        tailored_line_parts.append(f"with emphasis on {highlight_skills}")
    if top_titles:
        tailored_line_parts.append(f"and evidence pulled from {top_titles}")
    tailored_line = " ".join(tailored_line_parts).strip() + "."
    if normalize_skill_text(tailored_line) not in seen:
        lines.append(tailored_line)

    if not lines:
        lines.append(f"Tailored for {target_label}.")
    return lines[:3]


def _selected_item_lines(selected_items: list[dict], selected_skill_names: list[str], max_bullets_per_item: int) -> list[str]:
    lines: list[str] = []
    for item in selected_items:
        title = _clean_resume_line(item.get("title", "") or "Relevant experience")
        item_type = str(item.get("type") or "work").replace("evidence:", "").replace("_", " ").title()
        header = title if item_type.lower() == "work" else f"{title} [{item_type}]"
        lines.append(header)
        bullets = [f"- {bullet}" for bullet in _summarize_item_for_resume(item, selected_skill_names, max_bullets_per_item)]
        if bullets:
            lines.extend(bullets[:max_bullets_per_item])
        else:
            summary = _clean_resume_line(_clean_evidence_text_for_resume(item.get("summary", "")))
            if summary:
                lines.append(f"- {summary}")
        if item.get("links"):
            links = [str(link).strip() for link in item.get("links", []) if str(link).strip()]
            if links:
                lines.append(f"- Links: {', '.join(links[:3])}")
        lines.append("")
    return [line for line in lines if line != ""]


def _build_sections_from_resume_template(
    template_sections: list[ResumeSection],
    target_label: str,
    selected_skill_names: list[str],
    selected_items: list[dict],
    max_bullets_per_item: int,
    preserve_template_content: bool = True,
    header_lines: list[str] | None = None,
) -> list[ResumeSection]:
    fallback_template_sections = _load_default_resume_template_sections(None, include_content=False)
    fallback_lines_by_title = {section.title: [line for line in section.lines if line.strip()] for section in fallback_template_sections}
    summary_section = next((section for section in template_sections if section.title == "Summary"), None)
    skills_section = next((section for section in template_sections if section.title == "Skills"), None)
    experience_items = [item for item in selected_items if not _is_project_like_item(item)]
    project_items = [item for item in selected_items if _is_project_like_item(item)]
    experience_lines = _selected_item_lines(experience_items or selected_items[:2], selected_skill_names, max_bullets_per_item)
    project_lines = _selected_item_lines(project_items or selected_items[2:] or selected_items[:2], selected_skill_names, max_bullets_per_item)
    targeted_highlights_lines = _selected_item_lines(selected_items, selected_skill_names, max_bullets_per_item)
    merged_skill_names = _dedupe_preserve_order(
        [*_extract_original_skills(skills_section), *selected_skill_names]
    )

    sections: list[ResumeSection] = []
    inserted_summary = False
    inserted_skills = False
    inserted_highlights = False

    for section in template_sections:
        base_lines = [line for line in section.lines if line.strip()] if preserve_template_content else []
        if section.title == "Header":
            sections.append(ResumeSection(title="Header", lines=header_lines or base_lines or fallback_lines_by_title.get("Header", [])))
            continue
        if section.title == "Summary":
            sections.append(
                ResumeSection(
                    title="Summary",
                    lines=_build_targeted_summary_lines(summary_section, target_label, selected_skill_names, selected_items),
                )
            )
            inserted_summary = True
            continue
        if section.title == "Skills":
            if merged_skill_names:
                sections.append(ResumeSection(title="Skills", lines=_chunk_skill_lines(merged_skill_names)))
            elif base_lines:
                sections.append(ResumeSection(title="Skills", lines=base_lines))
            inserted_skills = True
            continue
        if section.title == "Experience":
            sections.append(ResumeSection(title="Experience", lines=experience_lines or base_lines or fallback_lines_by_title.get("Experience", [])))
            inserted_highlights = inserted_highlights or bool(experience_lines)
            continue
        if section.title == "Projects":
            sections.append(ResumeSection(title="Projects", lines=project_lines or base_lines or fallback_lines_by_title.get("Projects", [])))
            inserted_highlights = inserted_highlights or bool(project_lines)
            continue

        if (
            targeted_highlights_lines
            and not inserted_highlights
            and section.title in {"Experience", "Projects"}
            and not [line for line in section.lines if line.strip()]
        ):
            sections.append(ResumeSection(title=section.title, lines=targeted_highlights_lines))
            inserted_highlights = True
            continue

        if (
            targeted_highlights_lines
            and not inserted_highlights
            and section.title in {"Experience", "Projects", "Education", "Certifications"}
        ):
            sections.append(ResumeSection(title="Targeted Highlights", lines=targeted_highlights_lines))
            inserted_highlights = True

        sections.append(ResumeSection(title=section.title, lines=base_lines or fallback_lines_by_title.get(section.title, [])))

    if not inserted_summary:
        sections.insert(
            1 if sections and sections[0].title == "Header" else 0,
            ResumeSection(
                title="Summary",
                lines=_build_targeted_summary_lines(summary_section, target_label, selected_skill_names, selected_items),
            ),
        )
    if not inserted_skills and merged_skill_names:
        insert_index = 2 if sections and sections[0].title == "Header" else 1
        sections.insert(min(insert_index, len(sections)), ResumeSection(title="Skills", lines=_chunk_skill_lines(merged_skill_names)))
    if targeted_highlights_lines and not inserted_highlights:
        insert_index = len(sections)
        for idx, section in enumerate(sections):
            if section.title in {"Education", "Certifications", "Awards"}:
                insert_index = idx
                break
        sections.insert(insert_index, ResumeSection(title="Targeted Highlights", lines=targeted_highlights_lines))
    elif targeted_highlights_lines and not any(section.title == "Targeted Highlights" for section in sections):
        # Keep a dedicated rewriteable highlights section even when Experience/Projects were already populated.
        insert_index = len(sections)
        for idx, section in enumerate(sections):
            if section.title in {"Education", "Certifications", "Awards"}:
                insert_index = idx
                break
        sections.insert(insert_index, ResumeSection(title="Targeted Highlights", lines=targeted_highlights_lines))

    required_sections = {
        "Header": header_lines or fallback_lines_by_title.get("Header", []),
        "Summary": _build_targeted_summary_lines(summary_section, target_label, selected_skill_names, selected_items),
        "Skills": _chunk_skill_lines(merged_skill_names) or fallback_lines_by_title.get("Skills", []),
        "Experience": experience_lines or fallback_lines_by_title.get("Experience", []),
        "Projects": project_lines or fallback_lines_by_title.get("Projects", []),
        "Education": fallback_lines_by_title.get("Education", []),
    }
    existing_titles = {section.title for section in sections}
    for title in ("Header", "Summary", "Skills", "Experience", "Projects", "Education"):
        if title in existing_titles:
            continue
        lines = required_sections.get(title) or (
            [f"{title} details available upon request."]
            if title == "Education"
            else [f"{title} information is being tailored from your selected profile data."]
        )
        insert_index = len(sections)
        if title == "Header":
            insert_index = 0
        elif title == "Summary":
            insert_index = 1 if sections and sections[0].title == "Header" else 0
        sections.insert(insert_index, ResumeSection(title=title, lines=lines))
        existing_titles.add(title)

    return [section for section in sections if section.lines]

def _dedupe_skill_ids_by_name(skill_ids: Iterable[str], skill_docs: dict[str, dict]) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()
    for skill_id in skill_ids:
        name = normalize_skill_text(_display_skill_name(skill_docs.get(skill_id), skill_id))
        if not name or name in seen:
            continue
        seen.add(name)
        ordered.append(skill_id)
    return ordered


def _dedupe_skill_ids_by_terms(skill_ids: Iterable[str], skill_docs: dict[str, dict]) -> list[str]:
    ordered: list[str] = []
    seen_terms: set[str] = set()
    for skill_id in skill_ids:
        terms = _skill_terms(skill_docs.get(skill_id))
        if not terms:
            continue
        if terms & seen_terms:
            continue
        seen_terms.update(terms)
        ordered.append(skill_id)
    return ordered

def _serialize_history(doc: dict) -> JobMatchHistoryEntryOut:
    analysis = doc.get("analysis") or {}
    return JobMatchHistoryEntryOut(
        id=oid_str(doc.get("_id")),
        job_id=oid_str(doc.get("job_id")),
        title=doc.get("title"),
        company=doc.get("company"),
        location=doc.get("location"),
        source_history_id=oid_str(doc.get("source_history_id")) if doc.get("source_history_id") else None,
        match_score=float(analysis.get("match_score") or 0),
        semantic_alignment_score=float(analysis.get("semantic_alignment_score") or 0),
        matched_skills=[str(value) for value in analysis.get("matched_skills", []) or []],
        missing_skills=[str(value) for value in analysis.get("missing_skills", []) or []],
        strength_areas=[str(value) for value in analysis.get("strength_areas", []) or []],
        related_skills=[str(value) for value in analysis.get("related_skills", []) or []],
        tailored_resume_id=oid_str(doc.get("tailored_resume_id")) if doc.get("tailored_resume_id") else None,
        created_at=doc.get("created_at"),
        updated_at=doc.get("updated_at"),
    )

def _serialize_tailored_resume(doc: dict) -> TailoredResumeOut:
    return TailoredResumeOut(
        id=oid_str(doc.get("_id")),
        user_id=oid_str(doc.get("user_id")),
        job_id=str(doc.get("job_id")) if doc.get("job_id") is not None else None,
        resume_snapshot_id=str(doc.get("resume_snapshot_id")) if doc.get("resume_snapshot_id") is not None else None,
        template_source=str(doc.get("template_source") or "") or None,
        template=str(doc.get("template") or "ats_v1"),
        selected_skill_ids=[oid_str(value) for value in (doc.get("selected_skill_ids") or [])],
        selected_item_ids=[oid_str(value) for value in (doc.get("selected_item_ids") or [])],
        retrieved_context=[RAGContextItem(**item) for item in (doc.get("retrieved_context") or [])],
        sections=[ResumeSection(**section) for section in (doc.get("sections") or [])],
        plain_text=str(doc.get("plain_text") or ""),
        created_at=doc.get("created_at"),
    )


def _serialize_tailored_resume_list_entry(doc: dict, job_doc: dict | None = None) -> TailoredResumeListEntryOut:
    job_doc = job_doc or {}
    raw_job_id = doc.get("job_id")
    return TailoredResumeListEntryOut(
        id=oid_str(doc.get("_id")),
        user_id=oid_str(doc.get("user_id")),
        job_id=oid_str(raw_job_id) if raw_job_id is not None else None,
        job_title=str(job_doc.get("title") or "").strip() or None,
        company=str(job_doc.get("company") or "").strip() or None,
        location=str(job_doc.get("location") or "").strip() or None,
        template=str(doc.get("template") or "ats_v1"),
        selected_skill_count=len(doc.get("selected_skill_ids") or []),
        selected_item_count=len(doc.get("selected_item_ids") or []),
        created_at=doc.get("created_at"),
    )


def _serialize_tailored_resume_detail(doc: dict, job_doc: dict | None = None) -> TailoredResumeDetailOut:
    base = _serialize_tailored_resume(doc)
    job_doc = job_doc or {}
    return TailoredResumeDetailOut(
        **base.model_dump(),
        job_title=str(job_doc.get("title") or "").strip() or None,
        company=str(job_doc.get("company") or "").strip() or None,
        location=str(job_doc.get("location") or "").strip() or None,
        selected_skill_count=len(doc.get("selected_skill_ids") or []),
        selected_item_count=len(doc.get("selected_item_ids") or []),
    )

def _render_plain_text(sections: list[ResumeSection]) -> str:
    lines: list[str] = []
    for sec in sections:
        lines.append(sec.title.upper())
        for ln in sec.lines:
            lines.append(ln)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


async def _persist_job_ingest_snapshot(
    db,
    *,
    user_id: str,
    title: str | None,
    company: str | None,
    location: str | None,
    text: str,
) -> dict:
    skills = await _load_skill_catalog(db)
    extracted = _match_skills(text, skills)
    keywords = _tokenize_keywords(text, extracted)
    now = now_utc()
    doc = {
        "user_id": to_object_id(user_id),
        "title": title,
        "company": company,
        "location": location,
        "text": text,
        "extracted_skills": [e.model_dump() for e in extracted[:200]],
        "keywords": keywords,
        "created_at": now,
    }
    res = await db["job_ingests"].insert_one(doc)
    doc["_id"] = res.inserted_id
    return doc

@router.post("/job/ingest", response_model=JobIngestOut)
async def ingest_job(payload: JobIngestIn, user=Depends(require_user)):
    db = get_db()
    user_id = _scoped_user_id(user, payload.user_id)
    doc = await _persist_job_ingest_snapshot(
        db,
        user_id=user_id,
        title=payload.title,
        company=payload.company,
        location=payload.location,
        text=payload.text,
    )
    extracted = [ExtractedSkill(**entry) for entry in (doc.get("extracted_skills") or [])]
    keywords = list(doc.get("keywords") or [])

    preview = payload.text.strip().replace("\n", " ")
    if len(preview) > 220:
        preview = preview[:220] + "..."

    return JobIngestOut(
        id=oid_str(doc["_id"]),
        user_id=user_id,
        title=payload.title,
        company=payload.company,
        location=payload.location,
        text_preview=preview,
        extracted_skills=extracted[:75],
        keywords=keywords,
        created_at=doc.get("created_at"),
    )

@router.post("/match", response_model=JobMatchOut)
async def match_job(payload: dict, user=Depends(require_user)):
    """
    Expects:
    {
        "user_id": "...",
        "job_id": "..."   # this is the JobIngestOut.id from /tailor/job/ingest
    }
    """
    db = get_db()
    user_id = _scoped_user_id(user, payload.get("user_id"))
    user_doc = await db["users"].find_one(ref_query("_id", user_id), {"ai_preferences": 1})
    ai_preferences = normalize_ai_preferences((user_doc or {}).get("ai_preferences"))

    try:
        job_oid = ObjectId(payload["job_id"])
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid job_id")

    job_doc = await db["job_ingests"].find_one({"_id": job_oid, **ref_query("user_id", user_id)})
    if not job_doc:
        raise HTTPException(status_code=404, detail="Job ingest not found for user_id")

    skills = await _load_skill_catalog(db)
    recomputed_extracted = [entry.model_dump() for entry in _match_skills(job_doc.get("text", ""), skills)[:200]]
    recomputed_keywords = _tokenize_keywords(job_doc.get("text", ""), [ExtractedSkill(**entry) for entry in recomputed_extracted])
    if recomputed_extracted != (job_doc.get("extracted_skills") or []) or recomputed_keywords != (job_doc.get("keywords") or []):
        await db["job_ingests"].update_one(
            {"_id": job_oid},
            {
                "$set": {
                    "extracted_skills": recomputed_extracted,
                    "keywords": recomputed_keywords,
                }
            },
        )
        job_doc["extracted_skills"] = recomputed_extracted
        job_doc["keywords"] = recomputed_keywords

    extracted = job_doc.get("extracted_skills") or []
    ignored_skill_names = _normalize_ignored_skill_names(payload.get("ignored_skill_names"))
    ignored_terms = {normalize_skill_text(value) for value in ignored_skill_names if normalize_skill_text(value)}
    added_from_missing_skills = [
        {
            "skill_id": str(entry.get("skill_id") or "").strip(),
            "skill_name": str(entry.get("skill_name") or "").strip(),
        }
        for entry in (payload.get("added_from_missing_skills") or [])
        if str(entry.get("skill_id") or "").strip() and str(entry.get("skill_name") or "").strip()
    ]
    persist_history = bool(payload.get("persist_history", True))
    history_id = str(payload.get("history_id") or "").strip()
    requested_resume_snapshot_id = str(payload.get("resume_snapshot_id") or "").strip() or None
    raw_extracted_skill_ids = _dedupe_preserve_order(str(e.get("skill_id")) for e in extracted if e.get("skill_id"))

    conf = await _load_profile_confirmation(db, user_id)
    user_skill_ids = {
        str(entry.get("skill_id"))
        for entry in (conf or {}).get("confirmed", [])
        if entry.get("skill_id") is not None
    }
    items = await _load_user_items(db, user_id)
    # Pull the most relevant evidence/resume chunks first, then let the existing job
    # analysis logic incorporate that grounded context into semantic examples and item hits.
    retrieved_context = await retrieve_rag_context(
        db,
        user_id=user_id,
        query_text=job_doc.get("text", ""),
        preferences=ai_preferences,
        limit=6,
        source_types=("evidence", "resume_snapshot"),
    )
    context_score_by_item_id = _attach_retrieved_context_to_items(items, retrieved_context)
    item_skill_ids = {
        str(skill_id)
        for item in items
        for skill_id in (item.get("skill_ids") or [])
        if skill_id is not None
    }
    all_skill_docs = await _load_skills_by_ids(db, set(raw_extracted_skill_ids) | user_skill_ids | item_skill_ids)
    filtered_extracted = [
        entry
        for entry in extracted
        if str(entry.get("skill_id") or "") in all_skill_docs
        and not _is_ignored_skill(
            str(entry.get("skill_id") or ""),
            ignored_terms,
            all_skill_docs,
            str(entry.get("skill_name") or ""),
        )
    ]
    filtered_extracted_ids = _dedupe_preserve_order(str(entry.get("skill_id")) for entry in filtered_extracted if entry.get("skill_id"))
    extracted_skill_ids = [skill_id for skill_id in filtered_extracted_ids if skill_id in all_skill_docs]
    extracted_skill_ids = _dedupe_skill_ids_by_name(extracted_skill_ids, all_skill_docs)
    extracted_skill_ids = _dedupe_skill_ids_by_terms(extracted_skill_ids, all_skill_docs)
    priority_extracted = [entry for entry in filtered_extracted if str(entry.get("skill_id") or "") in all_skill_docs]
    required_ids, preferred_ids = _classify_job_skill_priority(job_doc.get("text", ""), priority_extracted, all_skill_docs)
    required_ids = set(_dedupe_skill_ids_by_name(required_ids, all_skill_docs))
    required_ids = set(_dedupe_skill_ids_by_terms(required_ids, all_skill_docs))
    preferred_ids = set(_dedupe_skill_ids_by_name(preferred_ids, all_skill_docs)) - required_ids
    preferred_ids = set(_dedupe_skill_ids_by_terms(preferred_ids, all_skill_docs)) - required_ids
    if not required_ids:
        # Many pasted job postings do not preserve clean section formatting. When we cannot
        # reliably detect a dedicated required section, treat the extracted job skills as the
        # effective required set so coverage, matched skills, and missing skills remain coherent.
        required_ids = set(extracted_skill_ids) - preferred_ids
        if not required_ids:
            required_ids = set(extracted_skill_ids)
    confirmed_term_lookup = {
        term
        for skill_id in user_skill_ids
        for term in _skill_terms(all_skill_docs.get(skill_id))
    }
    added_skill_id_lookup = {
        entry["skill_id"]
        for entry in added_from_missing_skills
        if entry.get("skill_id")
    }
    added_skill_term_lookup = {
        normalize_skill_text(entry["skill_name"])
        for entry in added_from_missing_skills
        if normalize_skill_text(entry.get("skill_name") or "")
    }
    matched_ids = []
    for skill_id in extracted_skill_ids:
        if skill_id in user_skill_ids:
            matched_ids.append(skill_id)
            continue
        if skill_id in added_skill_id_lookup:
            matched_ids.append(skill_id)
            continue
        extracted_terms = _skill_terms(all_skill_docs.get(skill_id))
        if extracted_terms & added_skill_term_lookup:
            matched_ids.append(skill_id)
            continue
        if _skill_terms(all_skill_docs.get(skill_id)) & confirmed_term_lookup:
            matched_ids.append(skill_id)
    matched_ids = _dedupe_skill_ids_by_name(matched_ids, all_skill_docs)
    matched_ids = _dedupe_skill_ids_by_terms(matched_ids, all_skill_docs)
    matched_skill_id_set = set(matched_ids)
    missing_ids = [skill_id for skill_id in extracted_skill_ids if skill_id not in matched_skill_id_set]
    matched_required_ids = [skill_id for skill_id in matched_ids if skill_id in required_ids]
    matched_preferred_ids = [skill_id for skill_id in matched_ids if skill_id in preferred_ids]
    matched_term_lookup = {
        term
        for skill_id in matched_ids
        for term in _skill_terms(all_skill_docs.get(skill_id))
    }
    item_text = " ".join(_normalize_text_blob(item) for item in items)
    matched_item_hits: list[dict] = []
    evidence_backed_ids: set[str] = set()
    for item in items:
        item_terms = {
            term
            for sid in (item.get("skill_ids") or [])
            for term in _skill_terms(all_skill_docs.get(str(sid)))
        }
        if item_terms & matched_term_lookup:
            matched_item_hits.append(item)
            for skill_id in matched_ids:
                if _skill_terms(all_skill_docs.get(skill_id)) & item_terms:
                    evidence_backed_ids.add(skill_id)
        elif context_score_by_item_id.get(oid_str(item.get("_id")), 0.0) >= 0.28:
            matched_item_hits.append(item)
    keywords = [str(keyword or "").strip().lower() for keyword in (job_doc.get("keywords") or []) if str(keyword or "").strip()]
    keyword_overlap_terms = [keyword for keyword in keywords if keyword in item_text]
    keyword_overlap_count = len(keyword_overlap_terms)

    matched_docs = [all_skill_docs[skill_id] for skill_id in matched_ids if skill_id in all_skill_docs]

    semantic_alignment_score = 0.0
    personal_skill_vector_score = 0.0
    related_skill_ids: list[str] = []
    if job_doc.get("text"):
        ordered_user_skill_ids = [skill_id for skill_id in user_skill_ids if skill_id in all_skill_docs]
        skill_texts = [
            " ".join(
                [
                    all_skill_docs[skill_id].get("name", ""),
                    all_skill_docs[skill_id].get("category", ""),
                    " ".join(all_skill_docs[skill_id].get("aliases", []) or []),
                ]
            ).strip()
            for skill_id in ordered_user_skill_ids
        ]
        item_texts = [_normalize_text_blob(item) for item in items]
        vectors, _provider = await embed_texts([job_doc.get("text", "")] + skill_texts + item_texts, preferences=ai_preferences)
        if vectors:
            job_vec = vectors[0]
            skill_vectors = vectors[1 : 1 + len(skill_texts)]
            item_vectors = vectors[1 + len(skill_texts) :]
            ranked_related: list[tuple[float, str]] = []
            skill_similarity_by_id: dict[str, float] = {}
            for skill_id, skill_vec in zip(ordered_user_skill_ids, skill_vectors):
                sim = cosine_similarity(job_vec, skill_vec)
                skill_similarity_by_id[skill_id] = sim
                if skill_id in matched_skill_id_set:
                    continue
                if sim >= 0.18:
                    ranked_related.append((sim, skill_id))
            ranked_related.sort(key=lambda item: item[0], reverse=True)
            related_skill_ids = [skill_id for _score, skill_id in ranked_related[:5]]

            semantic_examples: list[tuple[float, str]] = []
            if item_vectors:
                extracted_skill_id_set = set(extracted_skill_ids)
                weighted_item_scores: list[float] = []
                for item, item_vec in zip(items, item_vectors):
                    sim = cosine_similarity(job_vec, item_vec)
                    item_skill_id_set = {str(sid) for sid in (item.get("skill_ids") or []) if sid is not None}
                    overlap_ratio = len(item_skill_id_set & extracted_skill_id_set) / max(1, len(item_skill_id_set | extracted_skill_id_set))
                    bonus = overlap_ratio * 0.18
                    item_terms = {
                        term
                        for sid in item_skill_id_set
                        for term in _skill_terms(all_skill_docs.get(sid))
                    }
                    if item_terms & matched_term_lookup:
                        bonus += 0.08
                    combined_score = sim + bonus
                    weighted_item_scores.append(combined_score)
                    title = str(item.get("title") or item.get("summary") or "Saved evidence").strip()
                    if combined_score >= 0.22 and title:
                        semantic_examples.append((combined_score, f"{title} aligns with the posting language and matched skills"))
                top_item_alignment = _top_average(weighted_item_scores, 5)
            else:
                top_item_alignment = 0.0

            matched_skill_alignment = _top_average(
                (skill_similarity_by_id.get(skill_id, 0.0) for skill_id in ordered_user_skill_ids if _skill_terms(all_skill_docs.get(skill_id)) & matched_term_lookup),
                5,
            )
            overall_skill_alignment = _top_average(skill_similarity_by_id.values(), 6)
            for skill_id in related_skill_ids[:3]:
                name = _display_skill_name(all_skill_docs.get(skill_id), skill_id)
                if name:
                    semantic_examples.append((skill_similarity_by_id.get(skill_id, 0.0), f"Related confirmed skill: {name}"))
            semantic_alignment_score = _clamp_score(
                (
                    (top_item_alignment * 0.50)
                    + (matched_skill_alignment * 0.35)
                    + (overall_skill_alignment * 0.15)
                ) * 100.0
            )
        else:
            semantic_examples = []
    else:
        semantic_examples = []

    if job_doc.get("text"):
        user_vector_doc = await _compute_and_store_user_skill_vector(db, user_id, ai_preferences=ai_preferences)
        user_vector = list(user_vector_doc.get("vector") or [])
        job_vectors, _provider = await embed_texts([job_doc.get("text", "")], preferences=ai_preferences)
        if user_vector and job_vectors and len(user_vector) == len(job_vectors[0]):
            personal_skill_vector_score = _clamp_score(cosine_similarity(user_vector, job_vectors[0]) * 100.0)

    required_coverage_score = _score_percent(len(matched_required_ids), len(required_ids)) if required_ids else 100.0
    preferred_coverage_score = _score_percent(len(matched_preferred_ids), len(preferred_ids)) if preferred_ids else 100.0
    overall_coverage_score = _score_percent(len(matched_ids), len(extracted_skill_ids))
    weighted_skill_score = _clamp_score(
        (required_coverage_score * 0.55)
        + (preferred_coverage_score * 0.20)
        + (overall_coverage_score * 0.25)
    )
    evidence_score = _score_percent(len(evidence_backed_ids), len(matched_ids) or len(extracted_skill_ids))
    evidence_gap_count = max(0, len(matched_ids) - len(evidence_backed_ids))
    keyword_score = _score_percent(keyword_overlap_count, len(keywords))
    overall_score = _clamp_score(
        (weighted_skill_score * 0.50)
        + (evidence_score * 0.22)
        + (keyword_score * 0.10)
        + (semantic_alignment_score * 0.18)
    )

    breakdown = [
        MatchScoreBreakdown(
            label="Required skill coverage",
            score=required_coverage_score,
            detail=(
                f"{len(matched_required_ids)} of {len(required_ids)} likely required skills are already confirmed"
                if required_ids
                else "No clearly required skill section was detected, so this score starts at full credit"
            ),
        ),
        MatchScoreBreakdown(
            label="Preferred and secondary coverage",
            score=_clamp_score((preferred_coverage_score * 0.6) + (overall_coverage_score * 0.4)),
            detail=(
                f"{len(matched_preferred_ids)} of {len(preferred_ids)} preferred skills are covered, with {len(matched_ids)} total matched job skills"
                if preferred_ids
                else f"{len(matched_ids)} of {len(extracted_skill_ids)} total extracted job skills are covered"
            ),
        ),
        MatchScoreBreakdown(
            label="Evidence support",
            score=evidence_score,
            detail=f"{len(evidence_backed_ids)} matched skills are backed by evidence, leaving {evidence_gap_count} matched skills without proof",
        ),
        MatchScoreBreakdown(
            label="Keyword overlap",
            score=keyword_score,
            detail=f"{keyword_overlap_count} of {len(keywords)} job keywords appear in your saved work history",
        ),
        MatchScoreBreakdown(
            label="Semantic alignment",
            score=semantic_alignment_score,
            detail="Concept-level similarity between the job description and your saved work, even when the exact words do not match",
        ),
        MatchScoreBreakdown(
            label="Personal skill vector",
            score=personal_skill_vector_score,
            detail="Embedding similarity between the job posting and your aggregated user profile built from confirmed skills, evidence, and resume context",
        ),
    ]

    missing_names = [_display_skill_name(all_skill_docs.get(skill_id), skill_id) for skill_id in missing_ids if skill_id in all_skill_docs]
    matched_names = [_display_skill_name(all_skill_docs.get(skill_id), skill_id) for skill_id in matched_ids if skill_id in all_skill_docs]
    related_skill_names = [all_skill_docs[skill_id].get("name", skill_id) for skill_id in related_skill_ids if skill_id in all_skill_docs]
    gap_reasoning_summary, gap_insights = _build_gap_insights(
        missing_ids,
        required_ids,
        preferred_ids,
        related_skill_ids,
        all_skill_docs,
    )
    semantic_alignment_examples = [text for _score, text in sorted(semantic_examples, key=lambda item: item[0], reverse=True)[:4]]
    if retrieved_context:
        for context in retrieved_context[:3]:
            title = str(context.get("title") or "Retrieved evidence").strip()
            snippet = str(context.get("snippet") or "").strip()
            if snippet:
                semantic_alignment_examples.append(f"{title}: {snippet[:140]}")
    semantic_alignment_examples = _dedupe_preserve_order(semantic_alignment_examples)[:4]
    strength_areas = _build_strength_areas(matched_docs, evidence_backed_ids, matched_item_hits)
    confidence_label = _match_confidence_label(overall_score)
    semantic_alignment_explanation = (
        "Semantic alignment estimates how closely your saved evidence, projects, and confirmed skills resemble the responsibilities and tools in the job posting, even when the wording is different."
    )
    personal_skill_vector_explanation = (
        "Your personal skill vector is a single embedding built from confirmed skills, evidence, and resume context, then compared directly against the job vector."
    )
    analysis_summary_parts = [
        f"{confidence_label} fit based on {len(matched_ids)} matched skills out of {len(extracted_skill_ids)} extracted job skills."
    ]
    if required_ids:
        analysis_summary_parts.append(f"You currently cover {len(matched_required_ids)} of {len(required_ids)} likely required skills.")
    if evidence_backed_ids:
        analysis_summary_parts.append(f"{len(evidence_backed_ids)} matched skills are already supported by evidence.")
    elif matched_ids:
        analysis_summary_parts.append("Your confirmed skills are not yet strongly supported by evidence.")
    analysis_summary = " ".join(analysis_summary_parts)

    next_steps: list[str] = []
    if missing_names:
        prefix = "Prioritize" if required_ids and any(skill_id in required_ids for skill_id in missing_ids) else "Add evidence or learning proof for"
        next_steps.append(f"{prefix} {', '.join(missing_names[:3])}")
    if evidence_gap_count > 0:
        next_steps.append("Attach evidence to more matched skills so the tailored resume can prove capability instead of only claiming it")
    if not matched_item_hits:
        next_steps.append("Add projects or evidence with measurable outcomes to improve resume tailoring")
    if related_skill_names:
        next_steps.append(f"Consider positioning related experience like {', '.join(related_skill_names[:2])} when exact matches are limited")

    result = JobMatchOut(
        job_id=oid_str(job_doc["_id"]),
        match_score=overall_score,
        match_confidence_label=confidence_label,
        analysis_summary=analysis_summary,
        resume_snapshot_id=requested_resume_snapshot_id,
        template_source="user_resume" if requested_resume_snapshot_id else "default_template",
        ignored_skill_names=ignored_skill_names,
        added_from_missing_skills=added_from_missing_skills,
        matched_skill_ids=matched_ids,
        matched_skills=matched_names,
        missing_skill_ids=missing_ids,
        missing_skills=missing_names,
        matched_skill_count=len(matched_names),
        missing_skill_count=len(missing_names),
        strength_areas=strength_areas,
        related_skills=related_skill_names,
        semantic_alignment_examples=semantic_alignment_examples,
        retrieved_context=[RAGContextItem(**context) for context in retrieved_context],
        gap_reasoning_summary=gap_reasoning_summary,
        gap_insights=gap_insights,
        score_breakdown=breakdown,
        recommended_next_steps=next_steps[:3],
        extracted_skill_count=len(extracted_skill_ids),
        confirmed_skill_count=len(user_skill_ids),
        required_skill_count=len(required_ids),
        required_matched_count=len(matched_required_ids),
        preferred_skill_count=len(preferred_ids),
        preferred_matched_count=len(matched_preferred_ids),
        evidence_aligned_count=len(evidence_backed_ids),
        evidence_gap_count=evidence_gap_count,
        keyword_overlap_count=keyword_overlap_count,
        keyword_overlap_terms=keyword_overlap_terms,
        semantic_alignment_score=semantic_alignment_score,
        semantic_alignment_explanation=semantic_alignment_explanation,
        personal_skill_vector_score=personal_skill_vector_score,
        personal_skill_vector_explanation=personal_skill_vector_explanation,
    )

    if not persist_history:
        return result

    now = now_utc()
    if history_id:
        try:
            history_oid = ObjectId(history_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid history_id")

        existing_history = await db["job_match_runs"].find_one({"_id": history_oid, **ref_query("user_id", user_id)}, {"job_id": 1})
        if not existing_history:
            raise HTTPException(status_code=404, detail="Saved job analysis not found")
        if oid_str(existing_history.get("job_id")) != oid_str(job_doc.get("_id")):
            raise HTTPException(status_code=400, detail="history_id does not belong to the provided job_id")

        result.history_id = history_id
        await db["job_match_runs"].update_one(
            {"_id": history_oid},
            {
                "$set": {
                    "title": job_doc.get("title"),
                    "company": job_doc.get("company"),
                    "location": job_doc.get("location"),
                    "text_preview": job_doc.get("text", "")[:220],
                    "analysis": result.model_dump(),
                    "updated_at": now,
                },
                "$unset": {"tailored_resume_id": ""},
            },
        )
        return result

    history_doc = {
        "user_id": to_object_id(user_id),
        "job_id": job_oid,
        "title": job_doc.get("title"),
        "company": job_doc.get("company"),
        "location": job_doc.get("location"),
        "text_preview": job_doc.get("text", "")[:220],
        "job_text_snapshot": job_doc.get("text", ""),
        "analysis": result.model_dump(),
        "created_at": now,
        "updated_at": now,
    }
    history_res = await db["job_match_runs"].insert_one(history_doc)
    result.history_id = oid_str(history_res.inserted_id)
    await db["job_match_runs"].update_one({"_id": history_res.inserted_id}, {"$set": {"analysis.history_id": result.history_id}})
    return result

@router.post("/preview", response_model=TailoredResumeOut)
async def preview_tailored_resume(payload: TailorPreviewIn, user=Depends(require_user)):
    db = get_db()
    user_id = _scoped_user_id(user, payload.user_id)
    user_doc = await db["users"].find_one(ref_query("_id", user_id), {"username": 1, "email": 1, "ai_preferences": 1})
    ai_preferences = normalize_ai_preferences((user_doc or {}).get("ai_preferences"))

    job_text = payload.job_text
    job_id = payload.job_id

    extracted: list[ExtractedSkill] = []
    keywords: list[str] = []

    if job_id:
        try:
            job_oid = ObjectId(job_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid job_id")

        job_doc = await db["job_ingests"].find_one({"_id": job_oid, **ref_query("user_id", user_id)})
        if not job_doc:
            raise HTTPException(status_code=404, detail="Job ingest not found for user_id")
        job_text = job_doc.get("text", "")
        extracted = [ExtractedSkill(**e) for e in (job_doc.get("extracted_skills") or [])]
        keywords = job_doc.get("keywords") or []
    else:
        if not job_text or len(job_text) < 50:
            raise HTTPException(status_code=400, detail="Provide job_id or job_text (>=50 chars)")
        skills = await _load_skill_catalog(db)
        extracted = _match_skills(job_text, skills)
        keywords = _tokenize_keywords(job_text, extracted)

    ignored_skill_names = _normalize_ignored_skill_names(payload.ignored_skill_names)
    ignored_terms = {normalize_skill_text(value) for value in ignored_skill_names if normalize_skill_text(value)}
    extracted_skill_ids_for_lookup = _dedupe_preserve_order(e.skill_id for e in extracted if e.skill_id)
    skill_docs_for_filter = await _load_skills_by_ids(db, extracted_skill_ids_for_lookup)
    extracted = [
        entry
        for entry in extracted
        if not _is_ignored_skill(entry.skill_id, ignored_terms, skill_docs_for_filter, entry.skill_name)
    ]

    ordered_job_skill_ids = _dedupe_preserve_order(e.skill_id for e in extracted[:50])
    job_skill_ids = set(ordered_job_skill_ids)
    keyword_set = set(keywords)

    # load user portfolio items and user evidence-backed items
    items = await _load_user_items(db, user_id)
    # Tailoring works better when item ranking can see which evidence chunks were
    # actually closest to the target role, not just which items happened to share tags.
    retrieved_context = await retrieve_rag_context(
        db,
        user_id=user_id,
        query_text=job_text or "",
        preferences=ai_preferences,
        limit=8,
        source_types=("evidence", "resume_snapshot"),
    )
    context_score_by_item_id = _attach_retrieved_context_to_items(items, retrieved_context)

    scored = []
    for it in items:
        bonus = context_score_by_item_id.get(oid_str(it.get("_id")), 0.0) * 6.0
        scored.append((_score_item(it, job_skill_ids, keyword_set) + bonus, it))
    scored.sort(key=lambda x: x[0], reverse=True)

    selected_items = [it for score, it in scored[: payload.max_items] if score > 0] or [it for score, it in scored[: payload.max_items]]
    selected_item_ids = [oid_str(it["_id"]) for it in selected_items if "_id" in it]

    # Select skills: prioritize skills that appear in both job and user's confirmed skills if available
    confirmed_skill_ids: set[str] = set()
    conf = await db["resume_skill_confirmations"].find_one({"user_id": {"$in": ref_values(user_id)}}, sort=[("created_at", -1)])
    if conf:
        for c in conf.get("confirmed", []) or []:
            sid = c.get("skill_id")
            if sid is None:
                continue
            confirmed_skill_ids.add(str(sid))

    selected_skill_ids = [sid for sid in ordered_job_skill_ids if sid in confirmed_skill_ids]
    if len(selected_skill_ids) < 10:
        # backfill by job skills
        for e in extracted:
            if e.skill_id not in selected_skill_ids:
                selected_skill_ids.append(e.skill_id)
            if len(selected_skill_ids) >= 15:
                break

    skill_docs = await _load_skills_by_ids(db, selected_skill_ids)
    skill_name_by_id: dict[str, str] = {
        skill_id: _display_skill_name(doc, skill_id)
        for skill_id, doc in skill_docs.items()
    }
    skill_names = _dedupe_preserve_order([skill_name_by_id.get(s, s) for s in selected_skill_ids if skill_name_by_id.get(s, s)])
    skill_line = ", ".join(skill_names)[:250]
    job_title = (job_doc.get("title") or "").strip() if job_id else ""
    company = (job_doc.get("company") or "").strip() if job_id else ""
    target_label = " ".join(part for part in [job_title, f"at {company}" if company else ""] if part).strip() or "this role"

    resume_snapshot = await _load_resume_template_snapshot(db, user_id, payload.resume_snapshot_id)
    if resume_snapshot:
        resume_raw_text = str(resume_snapshot.get("raw_text") or "")
        template_sections = _parse_resume_sections(resume_raw_text)
        template_source = "user_resume"
        preserve_template_content = True
        header_lines = _extract_resume_header_lines(resume_raw_text, user_doc)
    else:
        template_sections = _load_default_resume_template_sections(user_doc, include_content=False)
        template_source = "default_template" if template_sections else "generated_fallback"
        preserve_template_content = False
        header_lines = _build_default_header_lines(user_doc)

    if template_sections:
        sections = _build_sections_from_resume_template(
            template_sections=template_sections,
            target_label=target_label,
            selected_skill_names=skill_names,
            selected_items=selected_items,
            max_bullets_per_item=payload.max_bullets_per_item,
            preserve_template_content=preserve_template_content,
            header_lines=header_lines,
        )
    else:
        fallback_lines = [f"Tailored for {target_label}."]
        if skill_line:
            fallback_lines.append(f"Prioritized skills: {skill_line}.")
        sections = [
            ResumeSection(title="Summary", lines=fallback_lines),
        ]
        if skill_names:
            sections.append(ResumeSection(title="Skills", lines=_chunk_skill_lines(skill_names)))
        highlight_lines = _selected_item_lines(selected_items, payload.max_bullets_per_item)
        if highlight_lines:
            sections.append(ResumeSection(title="Targeted Highlights", lines=highlight_lines))
        sections.append(
            ResumeSection(
                title="Targeted Alignment",
                lines=[
                    f"Role focus: {target_label}",
                    f"Selected evidence items: {len(selected_items)}",
                ],
            )
        )

    # Store tailored resume record
    now = now_utc()
    record = {
        "user_id": to_object_id(user_id),
        "job_id": job_id,
        "resume_snapshot_id": resume_snapshot.get("_id") if resume_snapshot else None,
        "template_source": template_source,
        "job_text": job_text,
        "template": payload.template,
        "selected_skill_ids": selected_skill_ids,
        "selected_item_ids": selected_item_ids,
        "retrieved_context": retrieved_context,
        "sections": [s.model_dump() for s in sections],
        "plain_text": _render_plain_text(sections),
        "created_at": now,
        "updated_at": now,
    }
    res = await db["tailored_resumes"].insert_one(record)

    if job_id and ObjectId.is_valid(job_id):
        latest_history = await db["job_match_runs"].find_one(
            {"user_id": {"$in": ref_values(user_id)}, "job_id": ObjectId(job_id)},
            sort=[("created_at", -1)],
        )
        if latest_history:
            await db["job_match_runs"].update_one(
                {"_id": latest_history["_id"]},
                {
                    "$set": {
                        "tailored_resume_id": res.inserted_id,
                        "analysis.resume_snapshot_id": oid_str(resume_snapshot.get("_id")) if resume_snapshot else None,
                        "analysis.template_source": template_source,
                        "updated_at": now,
                    }
                },
            )

    return _serialize_tailored_resume({"_id": res.inserted_id, **record})


@router.get("/resumes", response_model=list[TailoredResumeListEntryOut])
async def list_tailored_resumes(user_id: str | None = None, limit: int = 100, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    capped_limit = max(1, min(limit, 1000))
    docs = await (
        db["tailored_resumes"]
        .find(ref_query("user_id", scoped_user_id))
        .sort("created_at", -1)
        .limit(capped_limit)
        .to_list(length=capped_limit)
    )

    job_ids: list[ObjectId] = []
    for doc in docs:
        raw_job_id = doc.get("job_id")
        if isinstance(raw_job_id, ObjectId):
            job_ids.append(raw_job_id)
        elif isinstance(raw_job_id, str) and ObjectId.is_valid(raw_job_id):
            job_ids.append(ObjectId(raw_job_id))

    job_docs = await db["job_ingests"].find(
        {"_id": {"$in": job_ids}},
        {"title": 1, "company": 1, "location": 1},
    ).to_list(length=len(job_ids) or 1)
    jobs_by_id = {oid_str(job.get("_id")): job for job in job_docs}

    return [
        _serialize_tailored_resume_list_entry(doc, jobs_by_id.get(oid_str(doc.get("job_id"))))
        for doc in docs
    ]


@router.delete("/resumes/{tailored_id}")
async def delete_tailored_resume(tailored_id: str, user_id: str | None = None, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    doc = await db["tailored_resumes"].find_one({"_id": oid, **ref_query("user_id", scoped_user_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    await db["tailored_resumes"].delete_one({"_id": oid})
    await db["job_match_runs"].update_many(
        {"user_id": {"$in": ref_values(scoped_user_id)}, "tailored_resume_id": oid},
        {"$unset": {"tailored_resume_id": ""}, "$set": {"updated_at": now_utc()}},
    )
    return {"ok": True, "id": tailored_id}


@router.get("/resumes/{tailored_id}", response_model=TailoredResumeDetailOut)
async def get_tailored_resume_detail(tailored_id: str, user_id: str | None = None, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    doc = await db["tailored_resumes"].find_one({"_id": oid, **ref_query("user_id", scoped_user_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    raw_job_id = doc.get("job_id")
    job_doc = None
    if isinstance(raw_job_id, ObjectId):
        job_doc = await db["job_ingests"].find_one({"_id": raw_job_id}, {"title": 1, "company": 1, "location": 1})
    elif isinstance(raw_job_id, str) and ObjectId.is_valid(raw_job_id):
        job_doc = await db["job_ingests"].find_one({"_id": ObjectId(raw_job_id)}, {"title": 1, "company": 1, "location": 1})

    return _serialize_tailored_resume_detail(doc, job_doc)

@router.get("/history", response_model=list[JobMatchHistoryEntryOut])
async def list_job_match_history(user_id: str | None = None, limit: int = 12, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    cursor = db["job_match_runs"].find(ref_query("user_id", scoped_user_id)).sort("created_at", -1).limit(max(1, min(limit, 30)))
    docs = await cursor.to_list(length=max(1, min(limit, 30)))
    return [_serialize_history(doc) for doc in docs]

@router.get("/history/{history_id}", response_model=JobMatchHistoryDetailOut | JobMatchCompareOut)
async def get_job_match_history_detail(history_id: str, user_id: str | None = None, left_id: str | None = None, right_id: str | None = None, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    if history_id == "compare":
        if not left_id or not right_id:
            raise HTTPException(status_code=400, detail="left_id and right_id are required")
        return await compare_job_match_history(user_id=scoped_user_id, left_id=left_id, right_id=right_id, user=user)
    try:
        oid = ObjectId(history_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid history id")

    doc = await db["job_match_runs"].find_one({"_id": oid, **ref_query("user_id", scoped_user_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="History entry not found")

    base = _serialize_history(doc)
    job_text = str(doc.get("job_text_snapshot") or "").strip() or None
    job_id = doc.get("job_id")
    if not job_text and job_id is not None and ObjectId.is_valid(str(job_id)):
        job_doc = await db["job_ingests"].find_one({"_id": ObjectId(str(job_id)), **ref_query("user_id", scoped_user_id)})
        if job_doc:
            job_text = job_doc.get("text")

    analysis = JobMatchOut(**(doc.get("analysis") or {}))
    return JobMatchHistoryDetailOut(
        **base.model_dump(),
        text_preview=doc.get("text_preview"),
        job_text=job_text,
        analysis=analysis,
    )


@router.post("/history/{history_id}/reanalyze", response_model=JobMatchOut)
async def reanalyze_job_match_history(history_id: str, user_id: str | None = None, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    try:
        history_oid = ObjectId(history_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid history id")

    history_doc = await db["job_match_runs"].find_one({"_id": history_oid, **ref_query("user_id", scoped_user_id)})
    if not history_doc:
        raise HTTPException(status_code=404, detail="History entry not found")

    source_text = str(history_doc.get("job_text_snapshot") or "").strip()
    if not source_text:
        job_id = history_doc.get("job_id")
        if job_id is not None and ObjectId.is_valid(str(job_id)):
            job_doc = await db["job_ingests"].find_one({"_id": ObjectId(str(job_id)), **ref_query("user_id", scoped_user_id)})
            source_text = str((job_doc or {}).get("text") or "").strip()
    if not source_text:
        raise HTTPException(status_code=400, detail="No saved job description is available for reanalysis")

    fresh_job_doc = await _persist_job_ingest_snapshot(
        db,
        user_id=scoped_user_id,
        title=str(history_doc.get("title") or "").strip() or None,
        company=str(history_doc.get("company") or "").strip() or None,
        location=str(history_doc.get("location") or "").strip() or None,
        text=source_text,
    )

    fresh_result = await match_job(
        {
            "job_id": oid_str(fresh_job_doc.get("_id")),
            "resume_snapshot_id": None,
            "persist_history": True,
        },
        user=user,
    )
    if fresh_result.history_id:
        await db["job_match_runs"].update_one(
            {"_id": ObjectId(fresh_result.history_id)},
            {
                "$set": {
                    "source_history_id": history_oid,
                    "job_text_snapshot": source_text,
                }
            },
        )
    return fresh_result

@router.delete("/history/{history_id}")
async def delete_job_match_history(history_id: str, user_id: str | None = None, user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    try:
        oid = ObjectId(history_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid history id")

    doc = await db["job_match_runs"].find_one({"_id": oid, **ref_query("user_id", scoped_user_id)}, {"title": 1, "company": 1})
    if not doc:
        raise HTTPException(status_code=404, detail="History entry not found")

    await db["job_match_runs"].delete_one({"_id": oid})
    return {
        "ok": True,
        "id": history_id,
        "title": doc.get("title") or doc.get("company") or "Saved job match",
    }

@router.get("/history/compare", response_model=JobMatchCompareOut)
async def compare_job_match_history(user_id: str | None = None, left_id: str = "", right_id: str = "", user=Depends(require_user)):
    db = get_db()
    scoped_user_id = _scoped_user_id(user, user_id)
    try:
        left_oid = ObjectId(left_id)
        right_oid = ObjectId(right_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid history id")

    docs = await db["job_match_runs"].find(
        {"_id": {"$in": [left_oid, right_oid]}, **ref_query("user_id", scoped_user_id)}
    ).to_list(length=2)
    by_id = {oid_str(doc["_id"]): doc for doc in docs}
    if left_id not in by_id or right_id not in by_id:
        raise HTTPException(status_code=404, detail="History entry not found")

    left = _serialize_history(by_id[left_id])
    right = _serialize_history(by_id[right_id])
    left_matched = set(left.matched_skills)
    right_matched = set(right.matched_skills)
    left_missing = set(left.missing_skills)
    right_missing = set(right.missing_skills)

    return JobMatchCompareOut(
        left=left,
        right=right,
        match_score_delta=round(right.match_score - left.match_score, 2),
        semantic_alignment_delta=round(right.semantic_alignment_score - left.semantic_alignment_score, 2),
        newly_matched_skills=sorted(right_matched - left_matched),
        newly_missing_skills=sorted(right_missing - left_missing),
        shared_strength_areas=sorted(set(left.strength_areas) & set(right.strength_areas)),
    )

@router.get("/settings/status", response_model=AISettingsStatusOut)
async def get_ai_settings_status():
    return AISettingsStatusOut(**get_inference_status())


@router.get("/settings/preferences", response_model=AISettingsDetailOut)
async def get_ai_settings_preferences(user=Depends(require_user)):
    return _build_ai_settings_detail(user)


@router.patch("/settings/preferences", response_model=AISettingsDetailOut)
async def patch_ai_settings_preferences(payload: AIPreferencesPatchIn, user=Depends(require_user)):
    db = get_db()
    merged = normalize_ai_preferences(
        {
            **((user or {}).get("ai_preferences") or {}),
            **payload.model_dump(exclude_none=True),
        }
    )
    await db["users"].update_one({"_id": user["_id"]}, {"$set": {"ai_preferences": merged}})
    updated = await db["users"].find_one({"_id": user["_id"]}, {"ai_preferences": 1})
    return _build_ai_settings_detail(updated)

@router.post("/{tailored_id}/rewrite", response_model=RewriteBulletsOut)
async def rewrite_tailored_resume_bullets(tailored_id: str, payload: RewriteBulletsIn):
    db = get_db()
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    doc = await db["tailored_resumes"].find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    sections = [ResumeSection(**section) for section in doc.get("sections", [])]
    relevant_index = next(
        (
            index
            for index, section in enumerate(sections)
            if section.title.lower() in {"relevant work", "targeted highlights"}
        ),
        None,
    )
    if relevant_index is None:
        relevant_index = next(
            (
                index
                for index, section in enumerate(sections)
                if section.title.lower() in {"experience", "projects"}
            ),
            None,
        )
    if relevant_index is None:
        raise HTTPException(status_code=400, detail="Tailored resume has no rewriteable work section")

    work_lines = sections[relevant_index].lines
    bullet_indices = [index for index, line in enumerate(work_lines) if str(line).startswith("- ")]
    bullets = [work_lines[index] for index in bullet_indices]
    if not bullets:
        raise HTTPException(status_code=400, detail="Tailored resume has no bullets to rewrite")

    job_text = str(doc.get("job_text") or "")
    if not job_text and doc.get("job_id") and ObjectId.is_valid(str(doc.get("job_id"))):
        job_doc = await db["job_ingests"].find_one({"_id": ObjectId(str(doc.get("job_id")))})
        job_text = str((job_doc or {}).get("text") or "")

    rewritten_bullets, provider = await rewrite_resume_bullets(job_text, bullets, payload.focus)
    rewritten_count = 0
    for line_index, bullet in zip(bullet_indices, rewritten_bullets):
        work_lines[line_index] = bullet
        rewritten_count += 1

    updated_sections = [section.model_dump() for section in sections]
    updated_at = now_utc()
    plain_text = _render_plain_text(sections)
    await db["tailored_resumes"].update_one(
        {"_id": oid},
        {
            "$set": {
                "sections": updated_sections,
                "plain_text": plain_text,
                "updated_at": updated_at,
                "rewrite_focus": payload.focus,
                "rewrite_provider": provider,
            }
        },
    )

    return RewriteBulletsOut(
        tailored_id=tailored_id,
        provider=provider,
        focus=payload.focus,
        rewritten_count=rewritten_count,
        sections=sections,
        plain_text=plain_text,
        updated_at=updated_at,
    )

def _docx_from_sections(sections: list[ResumeSection], out_path: str):
    from docx import Document
    doc = Document()
    for sec in sections:
        doc.add_heading(sec.title, level=2)
        for ln in sec.lines:
            if ln.startswith("- "):
                doc.add_paragraph(ln[2:], style="List Bullet")
            else:
                doc.add_paragraph(ln)
    doc.save(out_path)

def _pdf_from_sections(sections: list[ResumeSection], out_path: str):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.lib.utils import simpleSplit

    c = canvas.Canvas(out_path, pagesize=LETTER)
    width, height = LETTER
    x = 54
    y = height - 54
    max_width = width - (x * 2)
    line_h = 14

    def draw_wrapped(text: str, bold: bool = False, indent: int = 0):
        nonlocal y
        font_name = "Helvetica-Bold" if bold else "Helvetica"
        font_size = 12 if bold else 11
        wrapped = simpleSplit(text, font_name, font_size, max_width - indent)
        for line in wrapped or [""]:
            if y < 54:
                c.showPage()
                y = height - 54
            c.setFont(font_name, font_size)
            c.drawString(x + indent, y, line)
            y -= line_h

    for sec in sections:
        draw_wrapped(sec.title.upper(), bold=True)
        for ln in sec.lines:
            if ln.startswith("- "):
                bullet = u"\u2022"
                bullet_width = stringWidth(bullet, "Helvetica", 11)
                if y < 54:
                    c.showPage()
                    y = height - 54
                c.setFont("Helvetica", 11)
                c.drawString(x, y, bullet)
                draw_wrapped(ln[2:], indent=int(bullet_width + 10))
            else:
                draw_wrapped(ln)
        draw_wrapped("")

    c.save()

@router.get("/{tailored_id}/export/docx")
async def export_docx(tailored_id: str):
    db = get_db()
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    d = await db["tailored_resumes"].find_one({"_id": oid})
    if not d:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    sections = [ResumeSection(**s) for s in d.get("sections", [])]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    _docx_from_sections(sections, tmp.name)

    filename = f"tailored_resume_{tailored_id}.docx"
    return FileResponse(tmp.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

@router.get("/{tailored_id}/export/pdf")
async def export_pdf(tailored_id: str):
    db = get_db()
    try:
        oid = ObjectId(tailored_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid tailored_id")

    d = await db["tailored_resumes"].find_one({"_id": oid})
    if not d:
        raise HTTPException(status_code=404, detail="Tailored resume not found")

    sections = [ResumeSection(**s) for s in d.get("sections", [])]
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    tmp.close()
    _pdf_from_sections(sections, tmp.name)

    filename = f"tailored_resume_{tailored_id}.pdf"
    return FileResponse(tmp.name, media_type="application/pdf", filename=filename)
